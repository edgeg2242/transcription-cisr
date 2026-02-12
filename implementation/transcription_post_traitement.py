#!/usr/bin/env python3
"""
Workflow : Transcription Post-Traitement
Description : Transforme une transcription audio brute en document Word CISR conforme
"""

import os
import sys
import io
import logging
import argparse
import json
import re
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

# Fix encodage UTF-8 Windows
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# Imports spécifiques
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Import optionnel de docxcompose (fallback si non disponible)
try:
    from docxcompose.composer import Composer
    DOCXCOMPOSE_AVAILABLE = True
except ImportError:
    DOCXCOMPOSE_AVAILABLE = False

# Configuration
load_dotenv()
logging.basicConfig(
    level=os.getenv('LOG_LEVEL', 'INFO'),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class WorkflowError(Exception):
    """Exception personnalisée pour les erreurs de workflow."""
    pass


class TranscriptionData:
    """Conteneur pour les données de transcription."""
    def __init__(self):
        self.texte_brut = ""
        self.metadata = {}
        self.dictionnaire = {}
        self.texte_nettoye = ""
        self.texte_corrige = ""
        self.interventions = []  # Liste de {'locuteur': 'COMMISSAIRE', 'texte': '...'}
        self.paragraphes = []
        self.stats_corrections = []
        self.score_qa = {}


def charger_transcription_brute(chemin_txt):
    """
    Étape 1 : Charge la transcription brute depuis le fichier TXT.

    Args:
        chemin_txt (str): Chemin vers fichier transcription TXT

    Returns:
        str: Contenu texte brut

    Raises:
        WorkflowError: Si fichier introuvable ou illisible
    """
    try:
        logger.info("Étape 1 : Chargement transcription brute...")

        chemin = Path(chemin_txt)
        if not chemin.exists():
            raise WorkflowError(f"Fichier introuvable : {chemin_txt}")

        with open(chemin, 'r', encoding='utf-8') as f:
            contenu = f.read()

        logger.info(f"✅ Transcription chargée ({len(contenu)} caractères)")
        return contenu

    except Exception as e:
        raise WorkflowError(f"Échec chargement transcription : {e}") from e


def charger_metadata(chemin_json):
    """
    Étape 1b : Charge les métadonnées depuis le fichier JSON.

    Args:
        chemin_json (str): Chemin vers fichier metadata JSON

    Returns:
        dict: Métadonnées structurées

    Raises:
        WorkflowError: Si fichier introuvable ou JSON invalide
    """
    try:
        logger.info("Chargement métadonnées...")

        chemin = Path(chemin_json)
        if not chemin.exists():
            raise WorkflowError(f"Fichier metadata introuvable : {chemin_json}")

        with open(chemin, 'r', encoding='utf-8') as f:
            metadata = json.load(f)

        logger.info(f"✅ Métadonnées chargées (dossier: {metadata.get('dossier', 'N/A')})")
        return metadata

    except json.JSONDecodeError as e:
        raise WorkflowError(f"JSON invalide : {e}") from e
    except Exception as e:
        raise WorkflowError(f"Échec chargement metadata : {e}") from e


def charger_dictionnaire_corrections():
    """
    Étape 1c : Charge le dictionnaire de corrections linguistiques.

    Returns:
        dict: Dictionnaire de corrections

    Raises:
        WorkflowError: Si dictionnaire introuvable
    """
    try:
        logger.info("Chargement dictionnaire corrections...")

        chemin = Path(__file__).parent.parent / "Documentation" / "CISR_Corrections_Dictionary.json"

        if not chemin.exists():
            logger.warning(f"⚠️  Dictionnaire non trouvé : {chemin}")
            return {}

        with open(chemin, 'r', encoding='utf-8') as f:
            dictionnaire = json.load(f)

        nb_corrections = len(dictionnaire.get('corrections_linguistiques', {}))
        logger.info(f"✅ Dictionnaire chargé ({nb_corrections} corrections)")
        return dictionnaire

    except Exception as e:
        logger.warning(f"⚠️  Échec chargement dictionnaire : {e}")
        return {}


def extraire_texte_integral(contenu_brut):
    """
    Extrait la section "TEXTE INTEGRAL" de la transcription brute.

    Args:
        contenu_brut (str): Contenu complet du fichier TXT

    Returns:
        str: Texte intégral uniquement
    """
    # Chercher entre "TEXTE INTEGRAL" et "TRANSCRIPTION PAR LOCUTEUR"
    match = re.search(
        r'TEXTE INTEGRAL\s*={70}\s*\n(.*?)\n={70}\s*\nTRANSCRIPTION PAR LOCUTEUR',
        contenu_brut,
        re.DOTALL
    )

    if match:
        return match.group(1).strip()

    # Fallback : prendre tout le contenu
    logger.warning("⚠️  Section TEXTE INTEGRAL non trouvée, utilisation contenu complet")
    return contenu_brut


def extraire_interventions_par_locuteur(contenu_brut):
    """
    Extrait les interventions par locuteur depuis la section diarization.

    Args:
        contenu_brut (str): Contenu complet du fichier TXT

    Returns:
        dict: {locuteur: texte_concatene}
    """
    interventions = {}

    # Chercher toutes les sections LOCUTEUR X
    pattern = r'={70}\s*\nLOCUTEUR ([A-Z])\s*\n={70}\s*\n(.*?)(?=\n={70}|$)'
    matches = re.finditer(pattern, contenu_brut, re.DOTALL)

    for match in matches:
        locuteur = match.group(1)
        texte = match.group(2).strip()

        if locuteur not in interventions:
            interventions[locuteur] = []
        interventions[locuteur].append(texte)

    # Concaténer les interventions par locuteur
    for loc in interventions:
        interventions[loc] = " ".join(interventions[loc])

    return interventions


def nettoyer_texte(texte, dictionnaire):
    """
    Étape 2 : Nettoyage initial - Suppression éléments non substantiels.

    Args:
        texte (str): Texte brut à nettoyer
        dictionnaire (dict): Dictionnaire de corrections

    Returns:
        str: Texte nettoyé
    """
    try:
        logger.info("Étape 2 : Nettoyage texte...")

        texte_nettoye = texte
        elements_supprimes = dictionnaire.get('elements_a_supprimer', {})

        # Supprimer salutations ouverture
        for salutation in elements_supprimes.get('salutations_ouverture', []):
            texte_nettoye = texte_nettoye.replace(salutation, "")

        # Supprimer fragments isolés
        for fragment in elements_supprimes.get('fragments_isoles', []):
            texte_nettoye = texte_nettoye.replace(fragment, "")

        # Supprimer remerciements finaux
        for remerciement in elements_supprimes.get('remerciements_finaux', []):
            texte_nettoye = texte_nettoye.replace(remerciement, "")

        # Supprimer interactions procédurales
        for interaction in elements_supprimes.get('interactions_procedurales', []):
            texte_nettoye = texte_nettoye.replace(interaction, "")

        # Nettoyer espaces multiples
        texte_nettoye = re.sub(r'\s+', ' ', texte_nettoye)
        texte_nettoye = texte_nettoye.strip()

        logger.info(f"✅ Nettoyage terminé ({len(texte) - len(texte_nettoye)} caractères supprimés)")
        return texte_nettoye

    except Exception as e:
        raise WorkflowError(f"Échec nettoyage texte : {e}") from e


def appliquer_corrections_linguistiques(texte, dictionnaire):
    """
    Étape 3 : Corrections linguistiques automatiques.

    Args:
        texte (str): Texte à corriger
        dictionnaire (dict): Dictionnaire de corrections

    Returns:
        tuple: (texte_corrige, stats_corrections)
    """
    try:
        logger.info("Étape 3 : Application corrections linguistiques...")

        texte_corrige = texte
        stats = []

        corrections = dictionnaire.get('corrections_linguistiques', {})

        for original, info in corrections.items():
            if isinstance(info, dict):
                correction = info.get('correction', original)
                categorie = info.get('categorie', 'autre')
                contexte = info.get('contexte', '')
            else:
                correction = info
                categorie = 'autre'
                contexte = ''

            # Cas spécial : "soumis" → "sunnite" SEULEMENT dans contexte religieux
            if original == "soumis" and correction == "sunnite":
                # Remplacement contextuel : SEULEMENT après "musulman"
                pattern = r'\b(musulman|confession)\s+soumis\b'
                occurrences = len(re.findall(pattern, texte_corrige, re.IGNORECASE))
                if occurrences > 0:
                    texte_corrige = re.sub(pattern, r'\1 sunnite', texte_corrige, flags=re.IGNORECASE)
                    stats.append({
                        'type': categorie,
                        'original': original,
                        'correction': correction,
                        'occurrences': occurrences
                    })
                    logger.info(f"  ✓ {original} → {correction} ({occurrences}x) [contextuel]")
            else:
                # Remplacement global pour les autres corrections
                occurrences = texte_corrige.count(original)
                if occurrences > 0:
                    texte_corrige = texte_corrige.replace(original, correction)
                    stats.append({
                        'type': categorie,
                        'original': original,
                        'correction': correction,
                        'occurrences': occurrences
                    })
                    logger.info(f"  ✓ {original} → {correction} ({occurrences}x)")

        # Normaliser noms de famille en MAJUSCULES (heuristique simple)
        # Chercher pattern "Monsieur/Madame [Nom]" et mettre [Nom] en majuscules
        texte_corrige = re.sub(
            r'\b(Monsieur|Madame|M\.|Mme)\s+([A-Z][a-zàéèêëïôùûç]+)\b',
            lambda m: f"{m.group(1)} {m.group(2).upper()}",
            texte_corrige
        )

        logger.info(f"✅ Corrections appliquées ({len(stats)} types)")
        return texte_corrige, stats

    except Exception as e:
        raise WorkflowError(f"Échec corrections linguistiques : {e}") from e


def mapper_locuteurs(interventions_brutes, args):
    """
    Étape 4 : Mapping des locuteurs (A, B, C → rôles réels).

    Args:
        interventions_brutes (dict): {locuteur: texte}
        args: Arguments CLI avec mapping explicite

    Returns:
        dict: Mapping {locuteur_original: role_final}
    """
    try:
        logger.info("Étape 4 : Mapping des locuteurs...")

        mapping = {}

        # Heuristique : Locuteur avec intervention la plus longue = COMMISSAIRE
        if interventions_brutes:
            locuteur_principal = max(
                interventions_brutes.items(),
                key=lambda x: len(x[1])
            )[0]

            mapping[locuteur_principal] = "COMMISSAIRE"

            # Autres locuteurs
            autres = [loc for loc in interventions_brutes if loc != locuteur_principal]
            if autres:
                # Premier autre = demandeur (généralement)
                mapping[autres[0]] = f"M. {args.dossier.split('-')[-1] if args.dossier else 'DEMANDEUR'}"

                # Autres = avocat, interprète, etc.
                for i, loc in enumerate(autres[1:], start=1):
                    mapping[loc] = f"LOCUTEUR_{i+1}"

        for original, role in mapping.items():
            logger.info(f"  {original} → {role}")

        logger.info(f"✅ Mapping terminé ({len(mapping)} locuteurs)")
        return mapping

    except Exception as e:
        raise WorkflowError(f"Échec mapping locuteurs : {e}") from e


def structurer_dialogue(texte, mapping_locuteurs):
    """
    Étape 5 : Structuration en format dialogue avec préfixes locuteurs.

    Args:
        texte (str): Texte corrigé
        mapping_locuteurs (dict): Mapping locuteurs

    Returns:
        str: Texte structuré en dialogue
    """
    try:
        logger.info("Étape 5 : Structuration dialogue...")

        # Pour l'instant, on suppose que tout le texte est du COMMISSAIRE
        # (car dans l'exemple, c'est principalement la commissaire qui parle)

        role_principal = list(mapping_locuteurs.values())[0] if mapping_locuteurs else "COMMISSAIRE"

        texte_dialogue = f"{role_principal} : {texte}"

        logger.info(f"✅ Dialogue structuré")
        return texte_dialogue

    except Exception as e:
        raise WorkflowError(f"Échec structuration dialogue : {e}") from e


def creer_paragraphes(texte):
    """
    Étape 6 : Structuration en paragraphes logiques.

    PRIORITÉ 1 (CRITIQUE) : Implémentation détection de ruptures de paragraphes
    basée sur l'analyse du guide CISR et des documents manuels.

    Règles appliquées :
    1. Détection de marqueurs de transition (Donc, D'abord, Concernant, etc.)
    2. Détection de numérotation (Le premier/deuxième élément, etc.)
    3. Détection de changements thématiques
    4. Longueur maximum de ~500 caractères par paragraphe
    5. Préservation du format COMMISSAIRE : uniquement au premier paragraphe

    Args:
        texte (str): Texte en dialogue

    Returns:
        list: Liste de paragraphes
    """
    try:
        logger.info("Étape 6 : Création paragraphes...")

        # Marqueurs de rupture thématique (séparation paragraphes)
        # Basé sur analyse du document manuel NTEZIRYAYO-MC3-03924
        marqueurs_debut = [
            "Donc,",
            "D'abord,",
            "Concernant",
            "Ma décision",
            "Advenant",
            "Vous avez",
            "Vous êtes",
            "On dit",
            "Parmi",
            "J'ai également",
            "J'ai aussi",
            "Tout cela",
            "En ce qui concerne",
            "Pour les raisons",
            "Finalement,",
            "En conclusion,",
            "Une fois",
            "Si vous"
        ]

        # Patterns de numérotation
        patterns_numerotation = [
            r"^Le premier ",
            r"^Le deuxième ",
            r"^Le troisième ",
            r"^Le dernier ",
            r"^Un des derniers ",
        ]

        # PRIORITÉ 2 : Nettoyer les tags [A], [B], [C] avant traitement
        texte_nettoye = re.sub(r'\[([A-Z])\]\s*', '', texte)

        # Séparer le préfixe COMMISSAIRE du texte
        if texte_nettoye.startswith("COMMISSAIRE :"):
            prefixe = "COMMISSAIRE :"
            contenu = texte_nettoye[len(prefixe):].strip()
        elif "COMMISSAIRE :" in texte_nettoye:
            # Cas où COMMISSAIRE : apparaît plus tard
            prefixe = "COMMISSAIRE :"
            contenu = texte_nettoye.split("COMMISSAIRE :", 1)[1].strip()
        else:
            prefixe = ""
            contenu = texte_nettoye

        paragraphes = []
        lignes = contenu.split(". ")

        para_courant = ""

        for i, ligne in enumerate(lignes):
            ligne = ligne.strip()
            if not ligne:
                continue

            # Restaurer le point final si ce n'est pas la dernière ligne
            if i < len(lignes) - 1:
                ligne += "."

            # Vérifier si cette ligne commence un nouveau paragraphe
            debut_nouveau_para = False

            # Check 1: Marqueurs de début
            for marqueur in marqueurs_debut:
                if ligne.startswith(marqueur):
                    debut_nouveau_para = True
                    break

            # Check 2: Patterns de numérotation
            if not debut_nouveau_para:
                for pattern in patterns_numerotation:
                    if re.match(pattern, ligne):
                        debut_nouveau_para = True
                        break

            # Check 3: Longueur maximum (~500 caractères)
            if len(para_courant) > 500 and not debut_nouveau_para:
                debut_nouveau_para = True

            # Si nouveau paragraphe détecté et para_courant non vide, sauvegarder
            if debut_nouveau_para and para_courant:
                paragraphes.append(para_courant.strip())
                para_courant = ligne + " "
            else:
                para_courant += ligne + " "

        # Ajouter le dernier paragraphe
        if para_courant.strip():
            paragraphes.append(para_courant.strip())

        # Réajouter le préfixe seulement au premier paragraphe
        if paragraphes and prefixe:
            paragraphes[0] = f"{prefixe} {paragraphes[0]}"

        # Si aucun paragraphe créé, retourner texte original nettoyé
        if not paragraphes:
            paragraphes = [f"{prefixe} {contenu}" if prefixe else contenu]

        logger.info(f"✅ Paragraphes créés ({len(paragraphes)})")
        return paragraphes

    except Exception as e:
        raise WorkflowError(f"Échec création paragraphes : {e}") from e


def remplir_metadonnees_si(page_couverture_path, metadata, output_path):
    """
    Remplit les métadonnées dans une page couverture SI vierge.

    Pour les documents SI, la page couverture fournie par le client est vierge
    et nécessite le remplissage manuel des métadonnées depuis le Work Assignment Excel.

    Structure SI attendue:
    - Paragraphe 1: Numéro de dossier SI (ex: "No de dossier de la SI /ID File No.: TB4-12345")
    - Paragraphe 2: No ID client (ex: "No ID client / Client ID No.: 1234567890")
    - Table 0 (11 lignes):
        - Row 0: Entre / Ministre / Between
        - Row 1: et / / and
        - Row 2: Intéressé(e)(s) / [NOM] / Person(s) Concerned
        - Row 3: Date(s) de l'audience / [DATE] / Date(s) of Hearing
        - Row 4: Lieu de l'audience / [LIEU] / Place of Hearing
        - Row 5: Date de la décision / [DATE] / Date of Decision
        - Row 6: Tribunal / [COMMISSAIRE] / Panel
        - Row 7: Conseil(s) de l'intéressé(e) / [CONSEIL] / Counsel for Person(s) Concerned
        - Row 8: Représentant(e)(s) désigné(e)(s) / [REP] / Designated Representative(s)
        - Row 9: Conseil du (de la) ministre / [MINISTRE] / Counsel for the Minister
        - Row 10: Interprète / [INTERPRETE] / Interpreter

    Args:
        page_couverture_path (str): Chemin vers la page couverture SI vierge
        metadata (dict): Métadonnées du work order (depuis metadata_work_order.json)
        output_path (str): Chemin de sortie pour la page couverture remplie

    Returns:
        str: Chemin vers la page couverture remplie
    """
    try:
        logger.info("Remplissage des métadonnées dans la page couverture SI...")

        doc = Document(page_couverture_path)

        # === REMPLIR EN-TÊTE ===
        # Paragraphe 1: Numéro de dossier
        dossier_info = metadata.get('dossier', {})
        numero_dossier = dossier_info.get('numero', '')

        if len(doc.paragraphs) > 1:
            # Remplacer le placeholder dans le paragraphe du numéro de dossier
            for para in doc.paragraphs[:5]:
                if 'No de dossier' in para.text or 'ID File No' in para.text:
                    # Remplacer 000XXXXXXX0X par le vrai numéro
                    para.text = para.text.replace('000XXXXXXX0X', numero_dossier)
                    para.text = para.text.replace('0X ', numero_dossier + ' ')
                    logger.info(f"   ✓ Numéro dossier SI: {numero_dossier}")
                    break

        # Paragraphe 2: No ID client (IUC)
        iuc = dossier_info.get('iuc', '')
        if iuc:
            for para in doc.paragraphs[:5]:
                if 'No ID client' in para.text or 'Client ID No' in para.text:
                    para.text = para.text.replace('XXXXXXXX', iuc)
                    logger.info(f"   ✓ IUC/UCI: {iuc}")
                    break

        # === REMPLIR TABLEAU MÉTADONNÉES ===
        if len(doc.tables) > 0:
            table = doc.tables[0]

            participants = metadata.get('participants', {})
            audience = metadata.get('audience', {})

            # Mapping des lignes du tableau SI
            # Row 2: Intéressé(e)(s) - Demandeur
            if len(table.rows) > 2:
                demandeur = participants.get('demandeur', '')
                if demandeur:
                    table.rows[2].cells[1].text = demandeur
                    logger.info(f"   ✓ Intéressé(e): {demandeur}")

            # Row 3: Date(s) de l'audience
            if len(table.rows) > 3:
                date_audience = audience.get('date', '')
                if date_audience:
                    table.rows[3].cells[1].text = date_audience
                    logger.info(f"   ✓ Date audience: {date_audience}")

            # Row 4: Lieu de l'audience
            if len(table.rows) > 4:
                lieu = audience.get('lieu', '')
                if lieu:
                    table.rows[4].cells[1].text = lieu
                    logger.info(f"   ✓ Lieu: {lieu}")

            # Row 5: Date de la décision
            if len(table.rows) > 5:
                date_decision = audience.get('date_decision', '')
                if date_decision:
                    table.rows[5].cells[1].text = date_decision
                    logger.info(f"   ✓ Date décision: {date_decision}")

            # Row 6: Tribunal (Commissaire)
            if len(table.rows) > 6:
                commissaire = participants.get('commissaire', '')
                if commissaire:
                    table.rows[6].cells[1].text = commissaire
                    logger.info(f"   ✓ Tribunal: {commissaire}")

            # Row 7: Conseil(s) de l'intéressé(e)
            if len(table.rows) > 7:
                conseil = participants.get('conseil_demandeur', '')
                if conseil:
                    table.rows[7].cells[1].text = conseil
                    logger.info(f"   ✓ Conseil: {conseil}")

            # Row 8: Représentant(e)(s) désigné(e)(s)
            if len(table.rows) > 8:
                representant = participants.get('representant_designe', '')
                if representant:
                    table.rows[8].cells[1].text = representant
                    logger.info(f"   ✓ Représentant: {representant}")

            # Row 9: Conseil du (de la) ministre
            if len(table.rows) > 9:
                ministre = participants.get('conseil_ministre', '')
                if ministre:
                    table.rows[9].cells[1].text = ministre
                    logger.info(f"   ✓ Conseil ministre: {ministre}")

            # Row 10: Interprète
            if len(table.rows) > 10:
                interprete = participants.get('interprete', '')
                if interprete:
                    table.rows[10].cells[1].text = interprete
                    logger.info(f"   ✓ Interprète: {interprete}")

        # Sauvegarder la page couverture remplie
        doc.save(output_path)
        logger.info(f"✅ Page couverture SI remplie: {output_path}")

        return output_path

    except Exception as e:
        raise WorkflowError(f"Échec remplissage métadonnées SI: {e}") from e


def fusionner_page_couverture_et_contenu(page_couverture_path, contenu_path, output_path):
    """
    Fusionne la page couverture client avec le contenu de transcription.

    La page couverture client devient la première page du document final,
    suivie du contenu de la transcription.

    Utilise docxcompose si disponible, sinon une approche manuelle avec python-docx.

    Args:
        page_couverture_path (str): Chemin vers la page couverture (.docx)
        contenu_path (str): Chemin vers le document de contenu (.docx)
        output_path (str): Chemin de sortie pour le document fusionné

    Returns:
        str: Chemin vers le document fusionné
    """
    try:
        logger.info("Fusion de la page couverture avec le contenu...")

        if DOCXCOMPOSE_AVAILABLE:
            # Méthode 1: Utiliser docxcompose (meilleure qualité)
            logger.info("   Utilisation de docxcompose...")
            doc_couverture = Document(page_couverture_path)
            doc_contenu = Document(contenu_path)

            composer = Composer(doc_couverture)
            composer.append(doc_contenu)
            composer.save(output_path)
        else:
            # Méthode 2: Fusion manuelle avec python-docx
            # Copie le contenu du document de contenu après la page couverture
            logger.info("   docxcompose non disponible, utilisation de python-docx...")

            doc_couverture = Document(page_couverture_path)
            doc_contenu = Document(contenu_path)

            # Ajouter un saut de page après la page couverture
            from docx.enum.section import WD_SECTION
            from docx.oxml.ns import nsmap

            # Ajouter un saut de page
            doc_couverture.add_page_break()

            # Copier les paragraphes du document contenu
            for element in doc_contenu.element.body:
                # Copie profonde de l'élément XML
                new_element = copy.deepcopy(element)
                doc_couverture.element.body.append(new_element)

            # Sauvegarder
            doc_couverture.save(output_path)

        logger.info(f"✅ Document fusionné créé: {output_path}")
        return output_path

    except Exception as e:
        raise WorkflowError(f"Échec fusion documents: {e}") from e


def generer_document_word(data, args, output_path):
    """
    Étape 7 : Génération du document Word CISR.

    Args:
        data (TranscriptionData): Données complètes
        args: Arguments CLI
        output_path (Path): Chemin de sortie

    Returns:
        Path: Chemin du fichier généré
    """
    try:
        logger.info("Étape 7 : Génération document Word...")

        doc = Document()

        # Appliquer Arial 11pt au style Normal (police par défaut)
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Configuration marges CISR (asymétriques)
        section = doc.sections[0]
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(0.63)
        section.left_margin = Inches(0.50)
        section.right_margin = Inches(0.50)

        # EN-TÊTE CISR (alignement RIGHT, gras)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Dossier de la {args.section}")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run = p.add_run(f" / RPD File: {args.dossier}")
        run.font.name = 'Arial'
        run.font.size = Pt(11)

        if args.iuc:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run("IUC / UCI:")
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            # PRIORITÉ 3 : Support IUC multiples (séparés par virgule ou newline)
            # Format attendu: "1124759825" ou "1124759825,1124759991" ou "1124759825\n1124759991"
            iuc_valeurs = args.iuc.replace(',', '\n').strip()
            run = p.add_run(f" {iuc_valeurs}")
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # Ligne vierge
        doc.add_paragraph()

        if args.huis_clos:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run("Huis clos / Private Proceeding")
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # Lignes vierges
        doc.add_paragraph()
        doc.add_paragraph()

        # TITRE (PRIORITÉ 3 : Arial 11pt explicite)
        # Note: Utiliser add_paragraph au lieu de add_heading pour contrôle total formatage
        p_titre = doc.add_paragraph()
        run_titre = p_titre.add_run("TRANSCRIPTION DES Motifs et de la décision")
        run_titre.font.name = 'Arial'
        run_titre.font.size = Pt(11)
        p_titre.paragraph_format.space_after = Pt(12)

        # Ligne vierge
        doc.add_paragraph()

        # SOUS-TITRE (CENTER, BOLD, UNDERLINE)
        p = doc.add_paragraph()
        run = p.add_run("décision")
        run.bold = True
        run.underline = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ligne vierge
        doc.add_paragraph()

        # === SAUT DE SECTION : Changement marges pour corps texte ===
        # CRITIQUE #1 : Document manuel a 2 sections avec marges différentes
        # Section 1 (en-tête + tableaux) : Marges CISR (1.25/0.63/0.50/0.50)
        # Section 2 (corps texte) : Marges standard (0.89/1.00/1.00/1.00)
        from docx.enum.section import WD_SECTION

        # Insérer nouveau paragraphe vide puis saut de section
        doc.add_paragraph()
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)

        # Appliquer marges standard à Section 2 (corps texte)
        new_section.top_margin = Inches(0.89)
        new_section.bottom_margin = Inches(1.00)
        new_section.left_margin = Inches(1.00)
        new_section.right_margin = Inches(1.00)

        # CORPS (contenu) avec interligne et formatage
        # IMPORTANT #6 : Pass 7 - Import fonction détection titres
        import sys
        import os
        sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        from implementation.pass7_titres_sections_gras import detecter_titre_section

        for i, paragraphe in enumerate(data.paragraphes):
            # IMPORTANT #6 : Pass 7 - Détecter titres de sections pour formatage gras
            est_titre, nom_titre = detecter_titre_section(paragraphe)

            p = doc.add_paragraph(paragraphe)
            p.style = 'Normal'
            p.paragraph_format.line_spacing = 1.0

            # MINEUR #4 : Indentation première ligne pour locuteurs
            # Document manuel utilise indentation Word native au lieu d'espaces
            if paragraphe.strip().startswith(('COMMISSAIRE :', 'CONSEIL :', 'DEMANDEUR :', 'INTERPRÈTE :')):
                p.paragraph_format.first_line_indent = Inches(0.15)

            # Appliquer Arial 11pt à chaque run
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)

                # PRIORITÉ 2 : Premier paragraphe COMMISSAIRE en BOLD
                if i == 0 and paragraphe.startswith("COMMISSAIRE :"):
                    run.bold = True

                # IMPORTANT #6 : Appliquer GRAS aux titres de sections détectés
                if est_titre:
                    run.bold = True

            # Ajouter ligne vierge après (sauf dernier)
            if i < len(data.paragraphes) - 1:
                doc.add_paragraph()

        # === TABLEAUX MÉTADONNÉES CISR ===
        # Charger depuis metadata_work_order.json si disponible
        if hasattr(args, 'metadata_json') and args.metadata_json and os.path.exists(args.metadata_json):
            logger.info("Ajout des tableaux métadonnées CISR depuis metadata_work_order.json...")
            with open(args.metadata_json, 'r', encoding='utf-8') as f:
                metadata = json.load(f)

            # Ligne vierge
            doc.add_paragraph()

            # TABLEAU 1: Informations dossier (avec lignes vides intercalaires)
            # 9 lignes de données + 8 lignes vides = 17 lignes total
            table = doc.add_table(rows=17, cols=3)
            table.style = 'Table Grid'

            # IMPORTANT #2 : Définir largeurs colonnes explicitement (mesures document manuel)
            # Col 1 (FR) : 2.31 pouces, Col 2 (Valeur) : 3.00 pouces, Col 3 (EN) : 2.44 pouces
            table.columns[0].width = Inches(2.31)
            table.columns[1].width = Inches(3.00)
            table.columns[2].width = Inches(2.44)

            # Données bilingues FR/EN
            rows_data = [
                ("Demandeur(e)(s) d'asile", metadata.get('participants', {}).get('demandeur', ''), "Claimant(s)"),
                ("Date de l'audience", metadata.get('audience', {}).get('date', ''), "Date of hearing"),
                ("Lieu de l'audience", metadata.get('audience', {}).get('lieu', ''), "Place of hearing"),
                ("Date de la décision / et des motifs", metadata.get('audience', {}).get('date_decision', ''), "Date of decision / and reasons"),
                ("Tribunal", metadata.get('participants', {}).get('commissaire', ''), "Panel"),
                ("Conseil(s) du (de la/des) / demandeur(e)(s) d'asile", metadata.get('participants', {}).get('conseil_demandeur', ''), "Counsel(s) for the claimant(s)"),
                ("Représentant(e) désigné(e)", metadata.get('participants', {}).get('representant_designe', ''), "Designated representative"),
                ("Conseil du (de la) ministre", metadata.get('participants', {}).get('conseil_ministre', ''), "Counsel for the Minister"),
                ("Interprète", metadata.get('participants', {}).get('interprete', ''), "Interpreter"),
            ]

            # Remplir tableau avec lignes vides intercalaires (pattern: données, vide, données, vide, ...)
            row_idx = 0
            for i, (fr, value, en) in enumerate(rows_data):
                # Ligne de données
                table.rows[row_idx].cells[0].text = fr
                table.rows[row_idx].cells[1].text = str(value)
                table.rows[row_idx].cells[2].text = en

                # IMPORTANT #2 : Formatage professionnel avec gras + alignement
                # Col 1 (FR) : LEFT + GRAS
                # Col 2 (Valeur) : CENTER + Normal
                # Col 3 (EN) : RIGHT + GRAS
                for j, cell in enumerate(table.rows[row_idx].cells):
                    for paragraph in cell.paragraphs:
                        # Alignement selon colonne
                        if j == 0:  # Col 1 (FR)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif j == 1:  # Col 2 (Valeur)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif j == 2:  # Col 3 (EN)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            # Gras pour colonnes FR et EN (pas pour valeur)
                            if j in [0, 2]:  # Col 1 (FR) ou Col 3 (EN)
                                run.bold = True

                row_idx += 1

                # Ajouter ligne vide après (sauf pour la dernière ligne)
                if i < len(rows_data) - 1:
                    # row_idx pointe maintenant sur la ligne vide, on la laisse vide
                    row_idx += 1

            # TABLEAU 2: Blanc (séparateur)
            doc.add_paragraph()
            table2 = doc.add_table(rows=1, cols=1)
            table2.style = 'Table Grid'

        # MARQUEUR FIN
        # Ligne vierge
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("----------FIN DES MOTIFS ----------")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ligne vierge
        doc.add_paragraph()

        # CERTIFICATION (No Spacing avec Arial 11pt)
        p = doc.add_paragraph()
        p.add_run(f"Je, {args.transcripteur}, déclare que cette transcription est exacte.")
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        p = doc.add_paragraph()
        date_cert = datetime.now().strftime("%d %B %Y")
        p.add_run(date_cert)
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        p = doc.add_paragraph()
        p.add_run(f"Agence de transcription: {args.agence}")
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

        # === GESTION PAGE COUVERTURE CLIENT ===
        # La première page du document final est toujours la page couverture fournie par le client
        page_couverture_path = getattr(args, 'page_couverture', None)

        if page_couverture_path and os.path.exists(page_couverture_path):
            import tempfile

            # Sauvegarder le contenu dans un fichier temporaire
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                contenu_temp_path = tmp.name
            doc.save(contenu_temp_path)

            # === SI: Remplir les métadonnées dans la page couverture vierge ===
            # Pour les documents SI, la page couverture client est vierge et nécessite
            # le remplissage des métadonnées depuis le Work Assignment Excel
            if args.section == 'SI':
                logger.info("📝 Type SI détecté: remplissage des métadonnées page couverture...")

                # Charger les métadonnées depuis metadata_work_order.json
                metadata_for_si = {}
                if hasattr(args, 'metadata_json') and args.metadata_json and os.path.exists(args.metadata_json):
                    with open(args.metadata_json, 'r', encoding='utf-8') as f:
                        metadata_for_si = json.load(f)

                # Créer une version remplie de la page couverture SI
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_couv:
                    page_couverture_remplie = tmp_couv.name

                remplir_metadonnees_si(page_couverture_path, metadata_for_si, page_couverture_remplie)
                page_couverture_path = page_couverture_remplie

            # Fusionner page couverture + contenu
            logger.info("📄 Fusion page couverture client avec contenu transcription...")
            fusionner_page_couverture_et_contenu(page_couverture_path, contenu_temp_path, output_path)

            # Nettoyer fichier temporaire
            try:
                os.unlink(contenu_temp_path)
                if args.section == 'SI' and 'page_couverture_remplie' in locals():
                    os.unlink(page_couverture_remplie)
            except Exception:
                pass

            logger.info(f"✅ Document Word généré avec page couverture: {output_path}")
        else:
            # Pas de page couverture fournie: sauvegarder directement
            doc.save(output_path)
            logger.info(f"✅ Document Word généré (sans page couverture): {output_path}")

        return output_path

    except Exception as e:
        raise WorkflowError(f"Échec génération Word : {e}") from e


def valider_qa(doc_path, data):
    """
    Étape 9 : Validation QA automatique.

    Args:
        doc_path (Path): Chemin vers document Word
        data (TranscriptionData): Données de transcription

    Returns:
        dict: Résultats validation QA
    """
    try:
        logger.info("Étape 9 : Validation QA...")

        resultats = {
            'criteres_pass': [],
            'criteres_fail': [],
            'criteres_na': [],
            'score': 0,
            'total': 20
        }

        # Critères automatisables basiques
        doc = Document(doc_path)
        texte_complet = "\n".join([p.text for p in doc.paragraphs])

        # Critère 1 : En-tête présent
        if "Dossier de la" in texte_complet and "RPD File" in texte_complet:
            resultats['criteres_pass'].append(1)
        else:
            resultats['criteres_fail'].append(1)

        # Critère 4 : Titre présent
        if "TRANSCRIPTION DES Motifs" in texte_complet:
            resultats['criteres_pass'].append(4)
        else:
            resultats['criteres_fail'].append(4)

        # Critère 8 : Pas d'erreurs connues (FESPOLA, etc.)
        erreurs_connues = ['FESPOLA', 'CISSPOLA', 'fiscalat']
        if not any(err in texte_complet for err in erreurs_connues):
            resultats['criteres_pass'].append(8)
        else:
            resultats['criteres_fail'].append(8)

        # Critère 11 : Certification présente
        if "déclare que cette transcription est exacte" in texte_complet:
            resultats['criteres_pass'].append(11)
        else:
            resultats['criteres_fail'].append(11)

        # Critère 15 : Marqueur FIN présent
        if "FIN DES MOTIFS" in texte_complet:
            resultats['criteres_pass'].append(15)
        else:
            resultats['criteres_fail'].append(15)

        # Score
        resultats['score'] = len(resultats['criteres_pass'])
        statut = "PASS" if resultats['score'] >= 4 else "FAIL"

        logger.info(f"✅ Validation QA terminée : {statut} ({resultats['score']}/5 critères validés)")
        return resultats

    except Exception as e:
        logger.warning(f"⚠️  Échec validation QA : {e}")
        return {'score': 0, 'total': 20, 'criteres_pass': [], 'criteres_fail': []}


def generer_rapport(data, output_dir):
    """
    Étape 10 : Génération du rapport de transformation.

    Args:
        data (TranscriptionData): Données complètes
        output_dir (Path): Dossier de sortie

    Returns:
        Path: Chemin du rapport JSON
    """
    try:
        logger.info("Étape 10 : Génération rapport...")

        timestamp = datetime.now().isoformat()
        rapport = {
            'timestamp': timestamp,
            'input_file': str(data.metadata.get('fichiers_generes', [''])[0]),
            'corrections_appliquees': len(data.stats_corrections),
            'details_corrections': data.stats_corrections,
            'validation_qa': data.score_qa
        }

        rapport_path = output_dir / f"post_traitement_rapport_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        with open(rapport_path, 'w', encoding='utf-8') as f:
            json.dump(rapport, f, indent=2, ensure_ascii=False)

        logger.info(f"✅ Rapport généré : {rapport_path}")
        return rapport_path

    except Exception as e:
        raise WorkflowError(f"Échec génération rapport : {e}") from e


def main():
    """Point d'entrée principal du workflow."""
    parser = argparse.ArgumentParser(
        description="Workflow 2 : Transcription Post-Traitement CISR"
    )

    # Arguments obligatoires
    parser.add_argument('--input', required=True, help='Chemin transcription TXT brute')
    parser.add_argument('--metadata', required=True, help='Chemin métadonnées JSON')
    parser.add_argument('--dossier', required=True, help='Numéro dossier (ex: TC5-07390)')
    parser.add_argument('--section', required=True, choices=['SPR', 'SAR', 'SI', 'SAI'], help='Section CISR')
    parser.add_argument('--transcripteur', required=True, help='Nom du transcripteur')
    parser.add_argument('--agence', required=True, help='Nom agence transcription')

    # Arguments optionnels
    parser.add_argument('--iuc', help='Numéro IUC/UCI (10 chiffres)')
    parser.add_argument('--commissaire', help='Nom de la commissaire')
    parser.add_argument('--huis-clos', action='store_true', help='Marquer comme huis clos')
    parser.add_argument('--metadata-json', help='Fichier metadata_work_order.json (workflow 0) pour tableaux CISR')
    parser.add_argument('--page-couverture', help='Chemin vers la page couverture client (.docx) à utiliser comme première page')
    parser.add_argument('--output-dir', help='Dossier de sortie (défaut: même que input)')
    parser.add_argument('--skip-qa', action='store_true', help='Sauter validation QA')
    parser.add_argument('--dry-run', action='store_true', help='Mode test (pas de sauvegarde)')

    args = parser.parse_args()

    try:
        logger.info("="*70)
        logger.info("Workflow 2 : Transcription Post-Traitement CISR")
        logger.info("="*70)

        # Initialiser conteneur données
        data = TranscriptionData()

        # Étape 1 : Chargement
        contenu_brut = charger_transcription_brute(args.input)
        data.metadata = charger_metadata(args.metadata)
        data.dictionnaire = charger_dictionnaire_corrections()

        # Extraire texte intégral
        data.texte_brut = extraire_texte_integral(contenu_brut)

        # Extraire interventions par locuteur
        interventions_brutes = extraire_interventions_par_locuteur(contenu_brut)

        # Étape 2 : Nettoyage
        data.texte_nettoye = nettoyer_texte(data.texte_brut, data.dictionnaire)

        # Étape 3 : Corrections linguistiques
        data.texte_corrige, data.stats_corrections = appliquer_corrections_linguistiques(
            data.texte_nettoye,
            data.dictionnaire
        )

        # Étape 4 : Mapping locuteurs
        mapping_locuteurs = mapper_locuteurs(interventions_brutes, args)

        # Étape 5 : Structuration dialogue
        texte_dialogue = structurer_dialogue(data.texte_corrige, mapping_locuteurs)

        # Étape 6 : Création paragraphes
        data.paragraphes = creer_paragraphes(texte_dialogue)

        # Déterminer dossier de sortie
        if args.output_dir:
            output_dir = Path(args.output_dir)
        else:
            output_dir = Path(args.input).parent

        output_dir.mkdir(parents=True, exist_ok=True)

        # Nom fichier final : [TRANSCRIPTEUR]-[DOSSIER]-[DATE]-[SECTION].docx
        date_str = datetime.now().strftime("%Y-%m-%d")
        nom_fichier = f"{args.transcripteur.upper()}-{args.dossier}-{date_str}-{args.section}.docx"
        output_path = output_dir / nom_fichier

        if args.dry_run:
            logger.info(f"🧪 Mode DRY-RUN : Fichier serait généré à {output_path}")
        else:
            # Étape 7 : Génération document Word
            doc_path = generer_document_word(data, args, output_path)

            # Étape 9 : Validation QA
            if not args.skip_qa:
                data.score_qa = valider_qa(doc_path, data)
            else:
                logger.info("⏭️  Validation QA sautée (--skip-qa)")
                data.score_qa = {'score': 0, 'total': 20, 'statut': 'SKIPPED'}

            # Étape 10 : Rapport
            rapport_path = generer_rapport(data, output_dir)

        logger.info("="*70)
        logger.info("✅ Workflow terminé avec succès")
        if not args.dry_run:
            logger.info(f"📄 Document final : {output_path}")
            logger.info(f"📊 Rapport : {rapport_path}")
            if not args.skip_qa:
                logger.info(f"✅ Score QA : {data.score_qa['score']}/{data.score_qa['total']}")
        logger.info("="*70)

        return 0

    except WorkflowError as e:
        logger.error(f"❌ Erreur de workflow : {e}")
        return 1

    except KeyboardInterrupt:
        logger.info("\n⚠️  Workflow interrompu par l'utilisateur")
        return 130

    except Exception as e:
        logger.error(f"❌ Erreur inattendue : {e}", exc_info=True)
        return 1


if __name__ == "__main__":
    sys.exit(main())
