#!/usr/bin/env python3
"""
Workflow 0: Préparation Work Order CISR

Décompresse et prépare les work orders ZIP reçus de la CISR.
Extrait les métadonnées depuis les pages couvertures et organise les fichiers.
"""

import sys
import io

# Fix encodage UTF-8 Windows
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import os
import re
import json
import zipfile
import shutil
import argparse
import logging
import glob
from pathlib import Path
from datetime import datetime
from docx import Document
from openpyxl import load_workbook

# Configuration logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class WorkOrderError(Exception):
    """Exception pour erreurs de traitement work order"""
    pass


class TypeTranscription:
    """Enum pour les 4 types de transcriptions CISR"""
    SPR = "SPR"  # Section Protection des Réfugiés
    SAR = "SAR"  # Section Appel des Réfugiés
    SI = "SI"    # Section de l'Immigration
    SAI = "SAI"  # Section Appel de l'Immigration


def decompresser_zip(zip_path, extract_dir):
    """
    Étape 1: Décompresser le ZIP work order.

    Args:
        zip_path: Chemin vers fichier ZIP
        extract_dir: Dossier de destination

    Returns:
        str: Chemin du dossier extrait

    Raises:
        WorkOrderError: Si ZIP invalide ou erreur extraction
    """
    try:
        # Créer dossier extraction si inexistant
        os.makedirs(extract_dir, exist_ok=True)

        # Extraire ZIP (avec protection Zip Slip)
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            # Valider chaque membre contre path traversal (Zip Slip)
            extract_dir_real = os.path.realpath(extract_dir)
            for member in zip_ref.namelist():
                member_path = os.path.realpath(os.path.join(extract_dir, member))
                if not member_path.startswith(extract_dir_real + os.sep) and member_path != extract_dir_real:
                    raise WorkOrderError(f"Chemin suspect dans ZIP (path traversal): {member}")
            zip_ref.extractall(extract_dir)

        # Compter fichiers extraits
        fichiers_extraits = []
        for root, dirs, files in os.walk(extract_dir):
            fichiers_extraits.extend(files)

        logger.info(f"✅ ZIP décompressé: {len(fichiers_extraits)} fichiers extraits")

        return extract_dir

    except zipfile.BadZipFile as e:
        raise WorkOrderError(f"Fichier ZIP invalide: {e}")
    except Exception as e:
        raise WorkOrderError(f"Erreur décompression ZIP: {e}")


def parser_excel_work_order(excel_path, numero_dossier=None):
    """
    Parse le fichier Excel Work Order pour extraire métadonnées critiques.

    Colonnes extraites:
    - Col C: Surname (nom demandeur)
    - Col H: Language of Hearing
    - Col M: Length of Audio
    - Col Q: Word Count
    - Col U: Name of Transcriber
    - Col V: Recording Unit Remarks (instructions découpage audio)

    Args:
        excel_path: Chemin vers fichier Excel work order
        numero_dossier: Numéro de dossier à chercher (ex: "MC3-03924") - optionnel

    Returns:
        dict: Métadonnées enrichies avec section "transcription"
              {
                  "transcription": {
                      "transcripteur": "Samantha Pitt",
                      "agence": "RegDeck",
                      "word_count_attendu": 1335,
                      "duree_audio_heures": 0.00625,
                      "recording_remarks": "commence à 1:46",
                      "audio_decoupage": {
                          "start_time_seconds": 106,
                          "applique": false
                      }
                  },
                  "participants": {
                      "demandeur_nom_famille": "HAMMOUD"
                  }
              }

    Raises:
        WorkOrderError: Si Excel invalide ou colonnes manquantes
    """
    try:
        logger.info(f"📖 Parsing Excel Work Order: {os.path.basename(excel_path)}")

        # Charger Excel
        wb = load_workbook(excel_path, data_only=True)

        # Première feuille (Work Order principal)
        ws = wb.active

        # Initialiser métadonnées
        metadata_excel = {
            "transcription": {
                "transcripteur": None,
                "agence": None,
                "word_count_attendu": None,
                "duree_audio_heures": None,
                "recording_remarks": None,
                "audio_decoupage": {
                    "start_time_seconds": None,
                    "applique": False
                }
            },
            "participants": {
                "demandeur_nom_famille": None
            },
            "validation": {
                "langue_audience": None
            }
        }

        # === EXTRACTION COLONNES ===
        # NOTE: Excel utilise indexation 1-based (A=1, B=2, ...)
        # Colonnes cibles:
        # C=3, H=8, M=13, Q=17, U=21, V=22

        # === ÉTAPE 1: Trouver ligne d'en-têtes ===
        # Chercher ligne contenant "Surname" ou "File Number"
        header_row = None
        for row_idx in range(1, 30):
            for col_idx in range(1, 10):
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value and 'surname' in str(cell_value).lower():
                    header_row = row_idx
                    break
            if header_row:
                break

        if not header_row:
            raise WorkOrderError("Ligne d'en-têtes non trouvée dans Excel (colonne 'Surname')")

        logger.info(f"   Ligne d'en-têtes détectée: {header_row}")

        # === ÉTAPE 2: Trouver ligne de données ===
        data_row = None

        if numero_dossier:
            # Chercher ligne contenant le numéro de dossier (col D)
            logger.info(f"   Recherche dossier: {numero_dossier}")
            for row_idx in range(header_row + 1, ws.max_row + 1):
                file_num = ws.cell(row=row_idx, column=4).value  # Col D
                if file_num and numero_dossier in str(file_num):
                    data_row = row_idx
                    logger.info(f"   Dossier trouvé à ligne: {data_row}")
                    break

            if not data_row:
                raise WorkOrderError(f"Dossier {numero_dossier} non trouvé dans Excel")
        else:
            # Sans numéro de dossier, prendre première ligne de données
            data_row = header_row + 1
            logger.info(f"   Utilisation première ligne de données: {data_row}")

        # === COL C: Surname (Nom demandeur) ===
        surname = ws.cell(row=data_row, column=3).value
        if surname:
            surname_clean = str(surname).strip().upper()
            metadata_excel["participants"]["demandeur_nom_famille"] = surname_clean
            logger.info(f"   ✅ Col C (Surname): {surname_clean}")

        # === COL H: Language of Hearing ===
        language = ws.cell(row=data_row, column=8).value
        if language:
            language_clean = str(language).strip()
            metadata_excel["validation"]["langue_audience"] = language_clean
            logger.info(f"   ✅ Col H (Language): {language_clean}")

        # === COL M: Length of Audio ===
        length_audio = ws.cell(row=data_row, column=13).value
        if length_audio:
            try:
                # Peut être un nombre (heures) ou texte "1:30:00"
                if isinstance(length_audio, (int, float)):
                    duree_heures = float(length_audio)
                else:
                    # Parser format "HH:MM:SS" ou "MM:SS"
                    parts = str(length_audio).split(':')
                    if len(parts) == 3:
                        h, m, s = map(int, parts)
                        duree_heures = h + m/60 + s/3600
                    elif len(parts) == 2:
                        m, s = map(int, parts)
                        duree_heures = m/60 + s/3600
                    else:
                        duree_heures = None

                if duree_heures:
                    metadata_excel["transcription"]["duree_audio_heures"] = round(duree_heures, 5)
                    logger.info(f"   ✅ Col M (Length of Audio): {duree_heures:.4f} h")
            except Exception as e:
                logger.warning(f"   ⚠️  Col M parsing échoué: {e}")

        # === COL Q: Word Count ===
        word_count = ws.cell(row=data_row, column=17).value
        if word_count:
            try:
                wc_int = int(word_count)
                metadata_excel["transcription"]["word_count_attendu"] = wc_int
                logger.info(f"   ✅ Col Q (Word Count): {wc_int} mots")
            except Exception as e:
                logger.warning(f"   ⚠️  Col Q parsing échoué: {e}")

        # === COL U: Name of Transcriber ===
        transcriber = ws.cell(row=data_row, column=21).value
        if transcriber:
            transcriber_clean = str(transcriber).strip()
            metadata_excel["transcription"]["transcripteur"] = transcriber_clean

            # Extraire agence (souvent entre parenthèses)
            match_agence = re.search(r'\(([^)]+)\)', transcriber_clean)
            if match_agence:
                metadata_excel["transcription"]["agence"] = match_agence.group(1)

            logger.info(f"   ✅ Col U (Transcriber): {transcriber_clean}")

        # === COL V: Recording Unit Remarks (CRITIQUE) ===
        remarks = ws.cell(row=data_row, column=22).value
        if remarks:
            remarks_clean = str(remarks).strip()
            metadata_excel["transcription"]["recording_remarks"] = remarks_clean

            # Parser instructions de découpage
            # Exemples:
            # - "commence à 1:46"
            # - "commence à 0:33"
            # - "arrête à 8:30"
            match_start = re.search(r'commence à (\d+):(\d+)', remarks_clean, re.IGNORECASE)
            if match_start:
                minutes = int(match_start.group(1))
                seconds = int(match_start.group(2))
                start_seconds = minutes * 60 + seconds
                metadata_excel["transcription"]["audio_decoupage"]["start_time_seconds"] = start_seconds
                logger.info(f"   ✅ Col V (Recording Remarks): {remarks_clean}")
                logger.info(f"      🎯 Découpage détecté: démarrer à {start_seconds}s ({minutes}:{seconds:02d})")
            else:
                logger.info(f"   ✅ Col V (Recording Remarks): {remarks_clean}")
                logger.info(f"      ℹ️  Aucune instruction de découpage détectée")

        logger.info(f"✅ Excel Work Order parsé avec succès")

        return metadata_excel

    except FileNotFoundError:
        raise WorkOrderError(f"Fichier Excel non trouvé: {excel_path}")
    except Exception as e:
        raise WorkOrderError(f"Erreur parsing Excel: {e}")


def localiser_fichier_audio(dossier_path, numero_dossier):
    """
    Localiser fichier .a00 dans arborescence DAUDIO profonde (multi-Work Orders).

    Structure cible:
        MC3-56703/
        └── DAUDIO/
            └── 2025-12-09/              ← Date audience
                └── _0231/               ← Timestamp enregistrement (02h31)
                    └── MC3-56703/       ← Numéro dossier (répété)
                        └── 6703.a00     ← Fichier audio (4 derniers chiffres)

    Args:
        dossier_path: Chemin vers dossier MC3-xxxxx/
        numero_dossier: "MC3-56703"

    Returns:
        str: Chemin absolu vers fichier .a00

    Raises:
        FileNotFoundError: Si aucun fichier audio trouvé

    Exemple:
        >>> localiser_fichier_audio("extracted/MC3-56703", "MC3-56703")
        'extracted/MC3-56703/DAUDIO/2025-12-09/_0231/MC3-56703/6703.a00'
    """
    logger.info(f"   🔍 Recherche audio dans: {os.path.basename(dossier_path)}")

    # Recherche récursive de tous les .a00
    audio_files = glob.glob(f"{dossier_path}/**/DAUDIO/**/*.a00", recursive=True)

    if not audio_files:
        raise FileNotFoundError(f"Aucun fichier audio trouvé dans {dossier_path}")

    # Vérification cohérence nom fichier
    # Règle: MC3-56703 → 6703.a00 (4 derniers chiffres)
    dernier_4_chiffres = numero_dossier[-4:]  # "56703" → "6703"

    # CRITICAL: Pattern strict pour éviter matching incorrect
    # Chercher fichier qui commence EXACTEMENT par les 4 derniers chiffres
    # Ex: MC3-03924 doit matcher "3924.a00" PAS "3157.a00"
    fichiers_correspondants = []

    for audio in audio_files:
        basename = os.path.basename(audio)
        # Match strict: nom fichier doit commencer par dernier_4_chiffres
        if basename.startswith(dernier_4_chiffres + '.') or basename.startswith(dernier_4_chiffres + '_'):
            fichiers_correspondants.append(audio)

    if fichiers_correspondants:
        # Si plusieurs matches, prendre le premier
        audio_choisi = fichiers_correspondants[0]
        taille_mb = os.path.getsize(audio_choisi) / (1024 * 1024)
        logger.info(f"   ✅ Audio trouvé: {os.path.basename(audio_choisi)} ({taille_mb:.2f} MB)")

        if len(fichiers_correspondants) > 1:
            logger.warning(f"   ⚠️  Plusieurs fichiers audio trouvés pour {numero_dossier}:")
            for f in fichiers_correspondants:
                logger.warning(f"      - {os.path.basename(f)}")
            logger.warning(f"      Sélectionné: {os.path.basename(audio_choisi)}")

        return audio_choisi

    # ERREUR CRITIQUE: Aucun fichier ne correspond au numéro dossier
    logger.error(f"   ❌ ERREUR: Aucun fichier audio ne correspond au dossier {numero_dossier}")
    logger.error(f"      Attendu: {dernier_4_chiffres}.a00 ou {dernier_4_chiffres}_*.a00")
    logger.error(f"      Fichiers trouvés: {[os.path.basename(f) for f in audio_files]}")

    raise FileNotFoundError(
        f"Aucun fichier audio correspondant à {numero_dossier} (attendu: {dernier_4_chiffres}.a00). "
        f"Fichiers présents: {[os.path.basename(f) for f in audio_files]}"
    )


def lire_excel_tous_work_orders(excel_path):
    """
    Lire toutes les lignes de données Excel Work Order (ligne 14+).

    Structure Excel RCE-9878-AA:
        Lignes 1-2   : En-tête CISR bilingue (FR/EN)
        Lignes 3-10  : Métadonnées globales Work Order
        Ligne 13     : HEADERS (23 colonnes)
        Lignes 14-19 : DONNÉES (6 Work Orders, une ligne par dossier)

    Colonnes extraites:
        - Col 4 (D): File Number (MC3-xxxxx) - IDENTIFIANT UNIQUE
        - Col 6 (F): Date of Hearing
        - Col 13 (M): Length of Audio (HH:MM:SS)
        - Col 22 (V): Recording Unit Remarks (découpage audio)

    Args:
        excel_path: Chemin vers fichier Excel work order

    Returns:
        list[dict]: Liste de métadonnées pour chaque Work Order
            [
                {
                    'numero_dossier': 'MC3-56703',
                    'date_audience': '2025-12-09',
                    'duree_audio': '0:09:00',
                    'recording_remarks': 'commence à 1:46',
                    'ligne_excel': 14
                },
                ...
            ]

    Raises:
        WorkOrderError: Si Excel invalide ou structure incorrecte
    """
    logger.info(f"📖 Parsing Excel Work Order multi-dossiers: {os.path.basename(excel_path)}")

    try:
        # Charger Excel
        wb = load_workbook(excel_path, data_only=True)
        ws = wb.active

        # === ÉTAPE 1: Trouver ligne d'en-têtes ===
        header_row = None
        for row_idx in range(1, 30):
            for col_idx in range(1, 10):
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value and 'file number' in str(cell_value).lower():
                    header_row = row_idx
                    break
            if header_row:
                break

        if not header_row:
            raise WorkOrderError("Ligne d'en-têtes non trouvée (colonne 'File Number')")

        logger.info(f"   Ligne d'en-têtes détectée: {header_row}")

        # === ÉTAPE 2: Lire toutes les lignes de données ===
        work_orders_data = []
        row_idx = header_row + 1

        while row_idx <= ws.max_row:
            # Col D (4): File Number
            file_num = ws.cell(row=row_idx, column=4).value

            # Si ligne vide, arrêter lecture
            if not file_num or str(file_num).strip() == "":
                break

            # Extraire données ligne
            numero_dossier = str(file_num).strip()

            # Col F (6): Date of Hearing
            date_hearing = ws.cell(row=row_idx, column=6).value
            if date_hearing:
                # Convertir datetime vers string si nécessaire
                if hasattr(date_hearing, 'strftime'):
                    date_hearing = date_hearing.strftime('%Y-%m-%d')
                else:
                    date_hearing = str(date_hearing)

            # Col M (13): Length of Audio
            length_audio = ws.cell(row=row_idx, column=13).value
            duree_audio = None
            if length_audio:
                # Peut être "HH:MM:SS" ou nombre (heures)
                if isinstance(length_audio, (int, float)):
                    heures = float(length_audio)
                    hours = int(heures)
                    minutes = int((heures - hours) * 60)
                    duree_audio = f"{hours}:{minutes:02d}:00"
                else:
                    duree_audio = str(length_audio)

            # Col V (22): Recording Unit Remarks
            recording_remarks = ws.cell(row=row_idx, column=22).value
            if recording_remarks:
                recording_remarks = str(recording_remarks).strip()

            work_orders_data.append({
                'numero_dossier': numero_dossier,
                'date_audience': date_hearing,
                'duree_audio': duree_audio,
                'recording_remarks': recording_remarks,
                'ligne_excel': row_idx
            })

            logger.info(f"   ✅ Ligne {row_idx}: {numero_dossier} ({duree_audio})")

            row_idx += 1

        logger.info(f"📊 {len(work_orders_data)} Work Orders détectés dans Excel")

        return work_orders_data

    except Exception as e:
        raise WorkOrderError(f"Erreur parsing Excel multi-Work Orders: {e}")


def detecter_work_orders_multiples(extracted_dir, excel_path):
    """
    Détecter tous les Work Orders dans un ZIP décompressé (support multi-Work Orders).

    Algorithme:
        1. Lire Excel Work Order pour connaître nombre exact de dossiers
        2. Vérifier présence dossiers MC3-xxxxx correspondants
        3. Localiser pages couvertures (racine du ZIP)
        4. Localiser fichiers audio (dans chaque dossier DAUDIO)
        5. Associer: Excel ligne → dossier → page couverture → audio

    Args:
        extracted_dir: Chemin vers dossier ZIP décompressé
        excel_path: Chemin vers fichier Excel work order

    Returns:
        list[dict]: Liste de Work Orders détectés
            [
                {
                    'numero_dossier': 'MC3-16722',
                    'dossier_path': '/path/to/MC3-16722/',
                    'page_couverture': '/path/to/MC3-16722 SPR Page couverture.docx',
                    'audio_file': '/path/to/6722.a00',
                    'excel_work_order': '/path/to/Work Order RCE-9878-AA.xlsx',
                    'metadata_excel': {
                        'date_audience': '2025-12-09',
                        'duree_audio': '0:11:00',
                        'recording_remarks': None
                    }
                },
                ...
            ]

    Raises:
        WorkOrderError: Si incohérences détectées
    """
    logger.info("🔍 Détection Work Orders multiples...")

    # === ÉTAPE 1: Lire Excel pour référence ===
    work_orders_excel = lire_excel_tous_work_orders(excel_path)

    # === ÉTAPE 2: Détecter dossiers MC3-xxxxx ===
    dossiers_regex = re.compile(r'MC[2-5]-\d{5}')
    dossiers_trouves = []

    for item in os.listdir(extracted_dir):
        item_path = os.path.join(extracted_dir, item)
        if os.path.isdir(item_path) and dossiers_regex.match(item):
            dossiers_trouves.append(item)

    logger.info(f"📁 {len(dossiers_trouves)} dossiers MC détectés: {', '.join(sorted(dossiers_trouves))}")

    # === ÉTAPE 3: Valider cohérence Excel vs dossiers ===
    numeros_excel = set(wo['numero_dossier'] for wo in work_orders_excel)
    dossiers_set = set(dossiers_trouves)

    if numeros_excel != dossiers_set:
        manquants_excel = dossiers_set - numeros_excel
        manquants_dossiers = numeros_excel - dossiers_set

        if manquants_excel:
            logger.warning(f"⚠️  Dossiers présents mais absents Excel: {manquants_excel}")
        if manquants_dossiers:
            logger.warning(f"⚠️  Dossiers listés Excel mais absents: {manquants_dossiers}")

    # === ÉTAPE 4: Localiser pages couvertures ===
    pages_couvertures = {}
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                # Chercher pattern MC3-xxxxx dans nom fichier
                match = dossiers_regex.search(file)
                if match:
                    numero = match.group()
                    pages_couvertures[numero] = os.path.join(root, file)

    logger.info(f"📄 {len(pages_couvertures)} pages couvertures localisées")

    # === ÉTAPE 5: Assembler données ===
    work_orders = []

    for wo_excel in work_orders_excel:
        numero = wo_excel['numero_dossier']
        dossier_path = os.path.join(extracted_dir, numero)

        # Vérifier dossier existe
        if not os.path.isdir(dossier_path):
            logger.warning(f"⚠️  Dossier manquant: {numero}")
            continue

        # Page couverture
        page_couv = pages_couvertures.get(numero)
        if not page_couv:
            logger.warning(f"⚠️  Page couverture manquante pour {numero}")

        # Fichier audio
        audio_file = None
        try:
            audio_file = localiser_fichier_audio(dossier_path, numero)
        except FileNotFoundError as e:
            logger.warning(f"⚠️  {e}")

        work_orders.append({
            'numero_dossier': numero,
            'dossier_path': dossier_path,
            'page_couverture': page_couv,
            'audio_file': audio_file,
            'excel_work_order': excel_path,
            'metadata_excel': {
                'date_audience': wo_excel['date_audience'],
                'duree_audio': wo_excel['duree_audio'],
                'recording_remarks': wo_excel['recording_remarks']
            }
        })

    logger.info(f"✅ {len(work_orders)} Work Orders assemblés avec succès")

    # === VALIDATION FINALE ===
    erreurs = []
    for wo in work_orders:
        if not wo['page_couverture']:
            erreurs.append(f"{wo['numero_dossier']}: Page couverture manquante")
        if not wo['audio_file']:
            erreurs.append(f"{wo['numero_dossier']}: Fichier audio manquant")

    if erreurs:
        logger.warning("⚠️  Avertissements détection:")
        for err in erreurs:
            logger.warning(f"   - {err}")

    return work_orders


def identifier_type_transcription(extract_dir):
    """
    Étape 2: Identifier le type (SPR, SAR, SI, SAI).

    Stratégie:
    1. Chercher fichier Excel
    2. Analyser nom fichier (SAR/RAD, SPR/RPD, SI/ID, SAI/IAD)
    3. Fallback: Analyser structure dossiers

    Args:
        extract_dir: Dossier extrait du ZIP

    Returns:
        str: Type détecté ("SPR", "SAR", "SI", "SAI")
    """
    # Chercher fichier Excel
    excel_files = []
    for root, dirs, files in os.walk(extract_dir):
        for file in files:
            if file.endswith('.xlsx') and not file.startswith('~'):
                excel_files.append(os.path.join(root, file))

    if not excel_files:
        logger.warning("Aucun fichier Excel trouvé, défaut à SPR")
        return TypeTranscription.SPR

    excel_path = excel_files[0]
    filename = os.path.basename(excel_path).upper()
    fullpath = excel_path.upper()  # Utiliser chemin complet pour détecter SAR dans nom dossier

    # Détection par nom fichier ou chemin complet
    # Priorité 1: SPR/RPD (doit être vérifié AVANT SAR/RAD car "RAD" peut matcher "Refugies Protection Division")
    if "SPR" in fullpath or "RPD FILE" in fullpath or "BENCH" in fullpath:
        logger.info("✅ Type identifié: SPR (Section Protection des Réfugiés)")
        return TypeTranscription.SPR
    # Priorité 2: SAR/RAD
    elif "SAR" in fullpath or ("RAD" in fullpath and "PROTECTION" not in fullpath.upper()):
        logger.info("✅ Type identifié: SAR (Section Appel des Réfugiés)")
        return TypeTranscription.SAR
    # Priorité 3: SAI/IAD (doit être vérifié AVANT SI/ID car "SI" est sous-chaîne de "SAI")
    elif re.search(r'\bSAI\b', filename) or re.search(r'\bIAD\b', filename):
        logger.info("✅ Type identifié: SAI (Section Appel de l'Immigration)")
        return TypeTranscription.SAI
    # Priorité 4: SI/ID (word boundaries pour éviter faux positifs sur VERSION, DECISION, etc.)
    elif re.search(r'\bSI\b', filename) or re.search(r'\bID\b', filename):
        logger.info("✅ Type identifié: SI (Section de l'Immigration)")
        return TypeTranscription.SI

    # Défaut: SPR
    logger.warning(f"Type non détecté depuis nom fichier '{filename}', défaut à SPR")
    return TypeTranscription.SPR


def extraire_metadonnees_spr(page_couverture_path):
    """
    Étape 3: Extraire métadonnées depuis page couverture SPR.

    Structure attendue:
    - En-tête: Numéro dossier, IUC, Huis clos
    - Tableau 1 (9 lignes × 3 colonnes): Métadonnées bilingues FR/EN

    Args:
        page_couverture_path: Chemin vers fichier Word page couverture

    Returns:
        dict: Métadonnées extraites

    Raises:
        WorkOrderError: Si fichier invalide ou structure incorrecte
    """
    try:
        doc = Document(page_couverture_path)

        # --- EXTRACTION EN-TÊTE ---

        if len(doc.paragraphs) < 2:
            raise WorkOrderError("Page couverture: pas assez de paragraphes")

        # Paragraphe 1: "Dossier de la SPR / RPD File: TC5-07390"
        para_dossier = doc.paragraphs[0].text
        match_dossier = re.search(r'[A-Z]{2}\d-\d+', para_dossier)
        if not match_dossier:
            raise WorkOrderError(f"Numéro dossier non trouvé dans: {para_dossier}")
        numero_dossier = match_dossier.group()

        # Paragraphe 2: "IUC / UCI: 1118522122"
        para_iuc = doc.paragraphs[1].text
        match_iuc = re.search(r'\d{10}', para_iuc)
        iuc = match_iuc.group() if match_iuc else None

        # Chercher "Huis clos" dans les 5 premiers paragraphes
        huis_clos = False
        for i in range(min(5, len(doc.paragraphs))):
            if "Huis clos" in doc.paragraphs[i].text or "Private Proceeding" in doc.paragraphs[i].text:
                huis_clos = True
                break

        # --- EXTRACTION TABLEAU MÉTADONNÉES ---

        if len(doc.tables) == 0:
            raise WorkOrderError("Page couverture: aucun tableau trouvé")

        table = doc.tables[0]  # Premier tableau

        # Extraire lignes non-vides du tableau (ignorer lignes vides intercalaires)
        lignes_donnees = []
        for row in table.rows:
            # Vérifier si la ligne contient des données (cellule du milieu non vide)
            texte_milieu = row.cells[1].text.strip()
            if texte_milieu:
                lignes_donnees.append(row)

        if len(lignes_donnees) < 9:
            raise WorkOrderError(f"Page couverture: pas assez de lignes de données ({len(lignes_donnees)}, attendu 9)")

        # Mapping lignes de données (index 0-8)
        metadonnees = {
            'numero_dossier': numero_dossier,
            'iuc': iuc,
            'huis_clos': huis_clos,
            'section': 'SPR',
            'demandeur': lignes_donnees[0].cells[1].text.strip(),
            'date_audience': lignes_donnees[1].cells[1].text.strip(),
            'lieu_audience': lignes_donnees[2].cells[1].text.strip(),
            'date_decision': lignes_donnees[3].cells[1].text.strip(),
            'commissaire': lignes_donnees[4].cells[1].text.strip(),
            'conseil_demandeur': lignes_donnees[5].cells[1].text.strip(),
            'representant_designe': lignes_donnees[6].cells[1].text.strip(),
            'conseil_ministre': lignes_donnees[7].cells[1].text.strip(),
            'interprete': lignes_donnees[8].cells[1].text.strip()
        }

        logger.info(f"✅ Métadonnées extraites: {len(metadonnees)} champs")

        return metadonnees

    except Exception as e:
        raise WorkOrderError(f"Erreur extraction métadonnées SPR: {e}")


def parser_dates_multiples(texte_dates):
    """
    Parse dates multiples depuis texte français.

    Exemples:
    - "4 juin 2025" → ['2025-06-04']
    - "4 juin 2025 et le 19 septembre 2025" → ['2025-06-04', '2025-09-19']

    Args:
        texte_dates: Texte contenant dates françaises

    Returns:
        list: Liste de dates au format YYYY-MM-DD
    """
    # Dictionnaire mois français → numéro
    mois_fr = {
        'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4, 'mai': 5, 'juin': 6,
        'juillet': 7, 'août': 8, 'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12,
        'janv': 1, 'févr': 2, 'avr': 4, 'juil': 7, 'sept': 9, 'oct': 10, 'nov': 11, 'déc': 12
    }

    # Regex pour dates françaises
    pattern = r'(\d{1,2})\s+([a-zûéè]+)\s+(\d{4})'
    matches = re.findall(pattern, texte_dates, re.IGNORECASE)

    dates = []
    for jour, mois_str, annee in matches:
        mois_str = mois_str.lower()
        if mois_str in mois_fr:
            mois = mois_fr[mois_str]
            date_str = f"{annee}-{mois:02d}-{int(jour):02d}"
            dates.append(date_str)
        else:
            logger.warning(f"Mois non reconnu: {mois_str}")

    return dates


def extraire_metadonnees_sar(extracted_dir):
    """
    Étape 3bis: Extraire métadonnées SAR (multi-dossiers).

    Étapes:
    1. Parser toutes les pages couvertures (3+)
    2. Scanner dossiers audio récursivement
    3. Fusionner données
    4. Validation croisée

    Args:
        extracted_dir: Dossier extrait du ZIP

    Returns:
        dict: Métadonnées SAR avec structure cases[]

    Raises:
        WorkOrderError: Si données manquantes ou incohérentes
    """
    metadata = {
        'work_order': {},
        'cases': []
    }

    # === 1. PARSER PAGES COUVERTURES ===

    cover_pages = {}
    docx_files = []
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                docx_files.append(os.path.join(root, file))

    logger.info(f"Pages couvertures trouvées: {len(docx_files)}")

    for docx_file in docx_files:
        if "Page couverture" in docx_file or "Transcript" in docx_file or "MC5-" in os.path.basename(docx_file):
            try:
                doc = Document(docx_file)

                # Extraire numéros SAR et SPR depuis:
                # 1. Nom du fichier (ex: MC5-34977 Irb 101.41 Page couverture...)
                # 2. En-tête du document
                filename = os.path.basename(docx_file)
                para_0 = doc.paragraphs[0].text if len(doc.paragraphs) > 0 else ""
                para_1 = doc.paragraphs[1].text if len(doc.paragraphs) > 1 else ""

                # SAR number (MC5-xxxxx) - chercher dans nom fichier ET contenu
                match_sar = re.search(r'MC5-\d+', filename + para_0 + para_1)
                if not match_sar:
                    logger.warning(f"Numéro SAR non trouvé dans {filename}, ignoré")
                    continue
                sar_number = match_sar.group()

                # Extraire métadonnées depuis tableau
                # NOTE: Pages couvertures SAR ont structure simplifiée (1 ligne, infos de base seulement)
                claimant = ""
                hearing_dates = []
                panel = ""
                counsel = ""
                uci = None
                spr_number = None

                if len(doc.tables) > 0:
                    table = doc.tables[0]

                    # Extraire SPR number et IUC depuis tableau
                    for row in table.rows:
                        for cell in row.cells:
                            # SPR number (MC2-xxxxx ou MC3-xxxxx)
                            if not spr_number:
                                match_spr = re.search(r'MC[2-3]-\d+', cell.text)
                                if match_spr:
                                    spr_number = match_spr.group()

                            # IUC
                            if not uci:
                                match_uci = re.search(r'\d{10}', cell.text)
                                if match_uci:
                                    uci = match_uci.group()

                # Fallback: chercher dans paragraphes
                if not spr_number:
                    match_spr = re.search(r'MC[2-3]-\d+', para_0 + para_1)
                    spr_number = match_spr.group() if match_spr else None

                # Extraire aussi depuis paragraphes
                if not uci:
                    for p in doc.paragraphs[:10]:
                        match_uci = re.search(r'\d{10}', p.text)
                        if match_uci:
                            uci = match_uci.group()
                            break

                # Pour SAR, on extrait surtout les numéros de dossiers
                # Les métadonnées complètes seront dans le document final transcrit

                cover_pages[sar_number] = {
                    'spr_file': spr_number,
                    'claimant': claimant,
                    'hearing_dates': hearing_dates,
                    'panel': panel,
                    'counsel': counsel,
                    'uci': uci,
                    'cover_page_file': os.path.basename(docx_file)
                }

                logger.info(f"  - {sar_number}: {claimant}, {len(hearing_dates)} audience(s)")

            except Exception as e:
                logger.warning(f"Erreur parsing {os.path.basename(docx_file)}: {e}")
                continue

    if not cover_pages:
        raise WorkOrderError("SAR: Aucune page couverture valide trouvée")

    # === 2. SCANNER DOSSIERS AUDIO ===

    audio_mappings = {}
    for item in os.listdir(extracted_dir):
        item_path = os.path.join(extracted_dir, item)
        if os.path.isdir(item_path) and item.startswith("MC"):
            folder_name = item  # MC2-27593
            daudio_path = os.path.join(item_path, "DAUDIO")

            if os.path.exists(daudio_path):
                dates_times = {}
                for date_dir in os.listdir(daudio_path):
                    date_path = os.path.join(daudio_path, date_dir)
                    if os.path.isdir(date_path):
                        # Chercher fichiers audio dans ce dossier et sous-dossiers
                        audio_files = []
                        for root, dirs, files in os.walk(date_path):
                            for file in files:
                                if file.endswith(('.a00', '.wav', '.mp3')):
                                    audio_files.append(file)
                        if audio_files:
                            dates_times[date_dir] = audio_files

                if dates_times:
                    audio_mappings[folder_name] = dates_times
                    logger.info(f"  Audio: {folder_name} → {len(dates_times)} dates")

    # === 3. FUSIONNER DONNÉES ===

    for sar_id, cover_data in cover_pages.items():
        spr_id = cover_data['spr_file']

        case = {
            'case_id': sar_id,
            'spr_file': spr_id,
            'claimant': cover_data['claimant'],
            'panel': cover_data['panel'],
            'counsel': cover_data['counsel'],
            'uci': cover_data['uci'],
            'hearing_dates': cover_data['hearing_dates'],
            'cover_page_file': cover_data['cover_page_file'],
            'audio_folder': spr_id,
            'audio_data': audio_mappings.get(spr_id, {})
        }

        metadata['cases'].append(case)

    # === 4. VALIDATION ===

    for case in metadata['cases']:
        if not case.get('claimant'):
            logger.warning(f"SAR {case['case_id']}: Demandeur manquant")

        if not case['audio_data']:
            logger.warning(f"SAR {case['case_id']}: Aucun dossier audio trouvé pour SPR {case['spr_file']}")

    logger.info(f"✅ {len(metadata['cases'])} cas SAR identifiés")

    return metadata


def trouver_fichier_audio(dossier_base):
    """
    Étape 4: Localiser le(s) fichier(s) audio récursivement.

    Extensions supportées: .a00, .wav, .mp3, .m4a

    Args:
        dossier_base: Dossier racine pour la recherche

    Returns:
        list: Liste de chemins vers fichiers audio

    Raises:
        WorkOrderError: Si aucun fichier audio trouvé
    """
    extensions_audio = ('.a00', '.wav', '.mp3', '.m4a')
    fichiers_audio = []

    for root, dirs, files in os.walk(dossier_base):
        for file in files:
            if file.lower().endswith(extensions_audio):
                fichiers_audio.append(os.path.join(root, file))

    if not fichiers_audio:
        raise WorkOrderError("Aucun fichier audio trouvé dans le work order")

    # Trier par taille (plus gros en premier)
    fichiers_audio.sort(key=lambda x: os.path.getsize(x), reverse=True)

    logger.info(f"✅ Fichier(s) audio trouvé(s): {len(fichiers_audio)}")
    for audio in fichiers_audio:
        taille_mb = os.path.getsize(audio) / (1024 * 1024)
        logger.info(f"   - {os.path.basename(audio)} ({taille_mb:.2f} MB)")

    return fichiers_audio


def creer_structure_projet_spr(output_dir, project_name, fichiers, metadata):
    """
    Étape 5: Créer la structure de dossier du projet SPR.

    Args:
        output_dir: Dossier parent (ex: Test_Demandes)
        project_name: Nom projet (ex: HAMMOUD-SPR-07390)
        fichiers: Dict {'audio': path, 'page_couverture': path, 'excel': path}
        metadata: Métadonnées extraites

    Returns:
        str: Chemin du dossier projet créé
    """
    # Générer nom projet si non fourni
    if not project_name:
        demandeur = metadata.get('demandeur', 'UNKNOWN').split()[0]
        numero = metadata['numero_dossier'].split('-')[1]
        project_name = f"{demandeur}-SPR-{numero}"

    # Créer dossier projet
    project_dir = os.path.join(output_dir, project_name)
    os.makedirs(project_dir, exist_ok=True)

    # Copier fichier audio
    if 'audio' in fichiers and fichiers['audio']:
        audio_dest = os.path.join(project_dir, os.path.basename(fichiers['audio']))
        shutil.copy2(fichiers['audio'], audio_dest)
        logger.info(f"✅ Fichier audio copié: {os.path.basename(audio_dest)}")

    # Copier page couverture
    if 'page_couverture' in fichiers and fichiers['page_couverture']:
        cover_dest = os.path.join(project_dir, 'page_couverture_original.docx')
        shutil.copy2(fichiers['page_couverture'], cover_dest)
        logger.info(f"✅ Page couverture copiée: page_couverture_original.docx")

    # Copier Excel (optionnel)
    if 'excel' in fichiers and fichiers['excel']:
        excel_dest = os.path.join(project_dir, 'work_order.xlsx')
        shutil.copy2(fichiers['excel'], excel_dest)

    # Copier ZIP original (bonne pratique)
    if 'zip_original' in fichiers and fichiers['zip_original']:
        zip_dest = os.path.join(project_dir, 'work_order_original.zip')
        shutil.copy2(fichiers['zip_original'], zip_dest)

    logger.info(f"✅ Structure de projet créée: {project_dir}")

    return project_dir


def creer_structure_projet_sar(output_dir, bon_commande, metadata_sar, extracted_dir):
    """
    Étape 5 SAR: Créer la structure de dossier du projet SAR (multi-cases).

    Args:
        output_dir: Dossier parent (ex: Test_Demandes)
        bon_commande: Numéro bon commande (ex: RCE-9439-DD)
        metadata_sar: Métadonnées SAR avec cases[]
        extracted_dir: Dossier extrait du ZIP

    Returns:
        str: Chemin du dossier projet créé
    """
    # Créer dossier principal
    project_dir = os.path.join(output_dir, f"{bon_commande}-SAR")
    os.makedirs(project_dir, exist_ok=True)

    # Créer sous-dossier pour chaque cas
    for case in metadata_sar['cases']:
        case_dir = os.path.join(project_dir, 'cases', case['case_id'])
        os.makedirs(case_dir, exist_ok=True)

        # Copier page couverture
        cover_src_name = case['cover_page_file']
        cover_src = None
        for root, dirs, files in os.walk(extracted_dir):
            if cover_src_name in files:
                cover_src = os.path.join(root, cover_src_name)
                break

        if cover_src:
            cover_dest = os.path.join(case_dir, f"page_couverture_{case['case_id']}.docx")
            shutil.copy2(cover_src, cover_dest)

        # Copier fichiers audio
        if case.get('audio_data'):
            audio_dir = os.path.join(case_dir, 'audio')
            os.makedirs(audio_dir, exist_ok=True)

            spr_folder = case['audio_folder']
            audio_folder_src = os.path.join(extracted_dir, spr_folder, 'DAUDIO')

            if os.path.exists(audio_folder_src):
                for date, audio_files in case['audio_data'].items():
                    for audio_file in audio_files:
                        # Chercher fichier audio récursivement
                        for root, dirs, files in os.walk(audio_folder_src):
                            if audio_file in files:
                                audio_src = os.path.join(root, audio_file)
                                # Renommer avec date
                                audio_dest_name = f"{date}_{audio_file}"
                                audio_dest = os.path.join(audio_dir, audio_dest_name)
                                shutil.copy2(audio_src, audio_dest)

    logger.info(f"✅ Structure de projet SAR créée: {project_dir}")
    logger.info(f"   {len(metadata_sar['cases'])} cas organisés")

    return project_dir


def generer_metadata_json_sar(metadata_sar, bon_commande, output_path):
    """
    Étape 6 SAR: Générer le fichier metadata.json SAR.

    Args:
        metadata_sar: Dict métadonnées SAR avec cases[]
        bon_commande: Numéro bon commande
        output_path: Chemin fichier JSON de sortie
    """
    metadata_json = {
        "work_order": {
            "number": bon_commande,
            "type": "SAR",
            "date_traitement": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "cases": metadata_sar['cases']
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(metadata_json, f, indent=2, ensure_ascii=False)

    taille_kb = os.path.getsize(output_path) / 1024
    logger.info(f"✅ metadata_work_order.json SAR généré ({taille_kb:.1f} KB)")


def generer_metadata_json(metadonnees, fichiers, output_path, metadata_excel=None):
    """
    Étape 6: Générer le fichier metadata.json.

    Args:
        metadonnees: Dict métadonnées extraites
        fichiers: Dict chemins fichiers
        output_path: Chemin fichier JSON de sortie
        metadata_excel: Dict métadonnées Excel (optionnel, enrichissement)
    """
    # Construire structure JSON SPR
    metadata_json = {
        "work_order": {
            "type": "SPR",
            "date_traitement": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        },
        "dossier": {
            "numero": metadonnees['numero_dossier'],
            "section": metadonnees['section'],
            "langue": "Français",  # TODO: détecter langue
            "huis_clos": metadonnees['huis_clos'],
            "iuc": metadonnees['iuc']
        },
        "participants": {
            "demandeur": metadonnees['demandeur'],
            "commissaire": metadonnees['commissaire'],
            "conseil_demandeur": metadonnees['conseil_demandeur'],
            "conseil_ministre": metadonnees['conseil_ministre'],
            "representant_designe": metadonnees['representant_designe'],
            "interprete": metadonnees['interprete']
        },
        "audience": {
            "date": metadonnees['date_audience'],
            "lieu": metadonnees['lieu_audience'],
            "date_decision": metadonnees['date_decision']
        },
        "fichiers": {
            "audio_original": os.path.basename(fichiers.get('audio', '')),
            "page_couverture": "page_couverture_original.docx",
            "work_order_excel": "work_order.xlsx" if 'excel' in fichiers else None
        }
    }

    # === ENRICHISSEMENT EXCEL (si disponible) ===
    if metadata_excel:
        # Fusionner sections transcription, participants, validation
        if "transcription" in metadata_excel:
            metadata_json["transcription"] = metadata_excel["transcription"]

        if "participants" in metadata_excel and metadata_excel["participants"].get("demandeur_nom_famille"):
            # Enrichir section participants existante
            metadata_json["participants"]["demandeur_nom_famille"] = metadata_excel["participants"]["demandeur_nom_famille"]

        if "validation" in metadata_excel and metadata_excel["validation"].get("langue_audience"):
            # Override langue détectée depuis Excel si disponible
            metadata_json["dossier"]["langue"] = metadata_excel["validation"]["langue_audience"]

        logger.info("   📊 Métadonnées enrichies depuis Excel Work Order")

    # Sauvegarder JSON
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(metadata_json, f, indent=2, ensure_ascii=False)

    taille_kb = os.path.getsize(output_path) / 1024
    logger.info(f"✅ metadata_work_order.json généré ({taille_kb:.1f} KB)")


def valider_metadonnees_spr(metadonnees):
    """
    Étape 7: Validation des métadonnées extraites.

    Vérifications:
    - Champs obligatoires présents
    - Numéro dossier au bon format
    - IUC 10 chiffres (si présent)

    Args:
        metadonnees: Dict métadonnées à valider

    Raises:
        WorkOrderError: Si validation échoue
    """
    champs_obligatoires = [
        'numero_dossier', 'demandeur', 'commissaire',
        'date_audience', 'lieu_audience'
    ]

    for champ in champs_obligatoires:
        if not metadonnees.get(champ):
            raise WorkOrderError(f"Champ obligatoire manquant: {champ}")

    # Valider format numéro dossier
    if not re.match(r'[A-Z]{2}\d-\d+', metadonnees['numero_dossier']):
        raise WorkOrderError(f"Format numéro dossier invalide: {metadonnees['numero_dossier']}")

    # Valider IUC (si présent)
    if metadonnees.get('iuc') and not re.match(r'\d{10}', metadonnees['iuc']):
        raise WorkOrderError(f"Format IUC invalide: {metadonnees['iuc']}")

    logger.info("✅ Validation métadonnées réussie")


def main():
    """Point d'entrée principal du workflow 0"""
    parser = argparse.ArgumentParser(
        description='Workflow 0: Préparation Work Order CISR',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples:
  # Work order SPR
  python preparation_work_order.py \\
    --zip-path "Demo/demo SPR/OCT242025 RCC6337 BENCH FR - REGDECK FR.zip" \\
    --output-dir "Test_Demandes" \\
    --project-name "HAMMOUD-SPR-07390"

  # Work order SAR (à venir)
  python preparation_work_order.py \\
    --zip-path "Demo/demo SAR/RCE-9439-DD - Regdeck - FR - SAR - Full - 10TAT - 10.17h.zip" \\
    --output-dir "Test_Demandes"
        """
    )
    parser.add_argument('--zip-path', required=True,
                        help='Chemin vers le fichier ZIP work order')
    parser.add_argument('--output-dir', default='Test_Demandes',
                        help='Dossier de sortie pour les projets (défaut: Test_Demandes)')
    parser.add_argument('--project-name',
                        help='Nom du projet (auto-généré si omis)')

    args = parser.parse_args()

    try:
        logger.info("=" * 60)
        logger.info("Démarrage Workflow 0: Préparation Work Order CISR")
        logger.info("=" * 60)

        # Vérifier ZIP existe
        if not os.path.exists(args.zip_path):
            raise WorkOrderError(f"Fichier ZIP non trouvé: {args.zip_path}")

        # Créer dossier temporaire extraction (à la racine du projet, pas dans Demo)
        temp_dir = os.path.join(os.getcwd(), 'extracted_temp')

        # === ÉTAPE 1: Décompression ===
        logger.info("\n[ÉTAPE 1] Décompression du ZIP...")
        extract_dir = decompresser_zip(args.zip_path, temp_dir)

        # === ÉTAPE 2: Identification Type ===
        logger.info("\n[ÉTAPE 2] Identification du type de transcription...")
        type_transcription = identifier_type_transcription(extract_dir)

        # === ÉTAPE 3: Extraction Métadonnées ===
        logger.info("\n[ÉTAPE 3] Extraction des métadonnées...")

        if type_transcription == TypeTranscription.SPR:
            # --- MODE SPR ---

            # Chercher fichier Excel pour détecter multi-Work Orders
            excel_files = []
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    if file.endswith('.xlsx') and not file.startswith('~'):
                        excel_files.append(os.path.join(root, file))

            # === DÉTECTION MULTI-WORK ORDERS ===
            work_orders_multiples = None
            if excel_files:
                try:
                    # Tentative détection multi-Work Orders
                    logger.info("\n[ÉTAPE 3.5] Détection Work Orders multiples...")
                    work_orders_multiples = detecter_work_orders_multiples(extract_dir, excel_files[0])

                    if len(work_orders_multiples) > 1:
                        logger.info(f"🎯 Mode MULTI-WORK ORDERS: {len(work_orders_multiples)} dossiers détectés")
                    else:
                        logger.info(f"📋 Mode WORK ORDER SIMPLE: 1 dossier")
                        work_orders_multiples = None  # Traiter comme simple si 1 seul
                except Exception as e:
                    logger.warning(f"⚠️  Détection multi-WO échouée: {e}")
                    logger.warning("   Workflow continuera en mode simple (1 dossier)")
                    work_orders_multiples = None

            # === BRANCHE 1: MULTI-WORK ORDERS ===
            if work_orders_multiples and len(work_orders_multiples) > 1:
                # Traiter chaque Work Order séparément
                logger.info(f"\n🔄 Traitement de {len(work_orders_multiples)} Work Orders...")

                for i, wo in enumerate(work_orders_multiples, 1):
                    logger.info(f"\n{'='*60}")
                    logger.info(f"Work Order {i}/{len(work_orders_multiples)}: {wo['numero_dossier']}")
                    logger.info(f"{'='*60}")

                    # Extraction métadonnées
                    metadonnees = extraire_metadonnees_spr(wo['page_couverture'])

                    # Parser Excel pour ce dossier spécifique
                    metadata_excel = parser_excel_work_order(
                        wo['excel_work_order'],
                        numero_dossier=wo['numero_dossier']
                    )

                    # Créer structure projet
                    fichiers = {
                        'audio': wo['audio_file'],
                        'page_couverture': wo['page_couverture'],
                        'excel': wo['excel_work_order'],
                        'zip_original': args.zip_path
                    }

                    project_dir = creer_structure_projet_spr(
                        args.output_dir,
                        None,  # Auto-générer nom projet
                        fichiers,
                        metadonnees
                    )

                    # Générer metadata.json
                    metadata_path = os.path.join(project_dir, 'metadata_work_order.json')
                    generer_metadata_json(metadonnees, fichiers, metadata_path, metadata_excel)

                    # Validation
                    valider_metadonnees_spr(metadonnees)

                    logger.info(f"✅ {wo['numero_dossier']}: Traité avec succès")
                    logger.info(f"   📁 Dossier: {project_dir}")

                logger.info(f"\n✅ {len(work_orders_multiples)} Work Orders traités avec succès")

            # === BRANCHE 2: WORK ORDER SIMPLE ===
            else:
                # Mode simple (comme avant)
                logger.info("\n📋 Mode Work Order Simple (1 dossier)")

                # Chercher page couverture
                page_couverture_files = []
                for root, dirs, files in os.walk(extract_dir):
                    for file in files:
                        if file.endswith('.docx') and not file.startswith('~'):
                            page_couverture_files.append(os.path.join(root, file))

                if not page_couverture_files:
                    raise WorkOrderError("Aucune page couverture (.docx) trouvée dans le ZIP")

                page_couverture_path = page_couverture_files[0]
                logger.info(f"Page couverture: {os.path.basename(page_couverture_path)}")

                metadonnees = extraire_metadonnees_spr(page_couverture_path)

                # === ÉTAPE 4: Localisation Audio ===
                logger.info("\n[ÉTAPE 4] Localisation du fichier audio...")
                fichiers_audio = trouver_fichier_audio(extract_dir)

                # === ÉTAPE 4.5: Parsing Excel Work Order (NOUVEAU - Sprint 0.1) ===
                logger.info("\n[ÉTAPE 4.5] Parsing Excel Work Order...")

                metadata_excel = None
                if excel_files:
                    try:
                        # Parser Excel avec numéro de dossier pour trouver la bonne ligne
                        metadata_excel = parser_excel_work_order(
                            excel_files[0],
                            numero_dossier=metadonnees['numero_dossier']
                        )
                    except WorkOrderError as e:
                        logger.warning(f"⚠️  Parsing Excel échoué: {e}")
                        logger.warning("   Workflow continuera sans enrichissement Excel")
                else:
                    logger.warning("⚠️  Aucun fichier Excel trouvé, enrichissement impossible")

                # === ÉTAPE 5: Création Structure Projet ===
                logger.info("\n[ÉTAPE 5] Création de la structure de projet...")

                fichiers = {
                    'audio': fichiers_audio[0] if fichiers_audio else None,
                    'page_couverture': page_couverture_path,
                    'excel': excel_files[0] if excel_files else None,
                    'zip_original': args.zip_path
                }

                project_dir = creer_structure_projet_spr(
                    args.output_dir,
                    args.project_name,
                    fichiers,
                    metadonnees
                )

                # === ÉTAPE 6: Génération metadata.json ===
                logger.info("\n[ÉTAPE 6] Génération metadata_work_order.json...")
                metadata_path = os.path.join(project_dir, 'metadata_work_order.json')
                generer_metadata_json(metadonnees, fichiers, metadata_path, metadata_excel)

                # === ÉTAPE 7: Validation ===
                logger.info("\n[ÉTAPE 7] Validation des métadonnées...")
                valider_metadonnees_spr(metadonnees)

        elif type_transcription == TypeTranscription.SAR:
            # --- MODE SAR (multi-cases) ---

            metadata_sar = extraire_metadonnees_sar(extract_dir)

            # === ÉTAPE 5 SAR: Création Structure Projet ===
            logger.info("\n[ÉTAPE 5] Création de la structure de projet SAR...")

            # Extraire bon de commande depuis nom ZIP
            zip_basename = os.path.basename(args.zip_path)
            # Ex: "RCE-9439-DD - Regdeck - FR - SAR - Full - 10TAT - 10.17h.zip"
            match_bon = re.search(r'([A-Z]+-\d+-[A-Z]+)', zip_basename)
            bon_commande = match_bon.group(1) if match_bon else "SAR-UNKNOWN"

            if args.project_name:
                bon_commande = args.project_name.replace('-SAR', '')

            project_dir = creer_structure_projet_sar(
                args.output_dir,
                bon_commande,
                metadata_sar,
                extract_dir
            )

            # === ÉTAPE 6 SAR: Génération metadata.json ===
            logger.info("\n[ÉTAPE 6] Génération metadata_work_order.json SAR...")
            metadata_path = os.path.join(project_dir, 'metadata_work_order.json')
            generer_metadata_json_sar(metadata_sar, bon_commande, metadata_path)

            # === ÉTAPE 7 SAR: Validation ===
            logger.info("\n[ÉTAPE 7] Validation des métadonnées SAR...")
            if not metadata_sar.get('cases') or len(metadata_sar['cases']) == 0:
                raise WorkOrderError("SAR: Aucun case valide trouvé")

            logger.info(f"✅ Validation métadonnées SAR réussie ({len(metadata_sar['cases'])} cases)")

        else:
            # SI, SAI à implémenter
            raise WorkOrderError(f"Type {type_transcription} pas encore supporté (MVP SPR et SAR seulement)")

        # === NETTOYAGE ===
        logger.info("\n[NETTOYAGE] Suppression dossier temporaire...")
        shutil.rmtree(temp_dir, ignore_errors=True)

        # === SUCCÈS ===
        logger.info("\n" + "=" * 60)
        logger.info("✅ Work order préparé avec succès!")
        logger.info(f"📁 Dossier projet: {project_dir}")
        logger.info(f"📄 Métadonnées: {metadata_path}")
        logger.info("\nPrêt pour workflow 1 (réception_preparation)")
        logger.info("=" * 60)

        return 0

    except WorkOrderError as e:
        logger.error(f"\n❌ Erreur workflow 0: {e}")
        return 1
    except Exception as e:
        logger.error(f"\n❌ Erreur inattendue: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
