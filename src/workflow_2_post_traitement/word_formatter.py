"""
Génération de documents Word conformes aux normes CISR.

Responsabilités :
- Configuration marges asymétriques par type (SPR/SAR/SI/SAI)
- En-tête CISR bilingue (dossier, IUC, huis clos)
- Corps transcription avec formatage locuteurs
- Tableaux métadonnées bilingues (17L x 3C SPR, structure inversée SAR)
- Certification + marqueur FIN DES MOTIFS
- Fusion page couverture client + contenu
"""

import os
import json
import copy
import logging
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.common.constants import MARGINS_SPR, MARGINS_SAR, MARGINS_SI, MARGINS_SAI
from src.common.exceptions import WorkflowError
from src.workflow_2_post_traitement.section_formatter import detecter_titre_section

# Import optionnel de docxcompose
try:
    from docxcompose.composer import Composer
    DOCXCOMPOSE_AVAILABLE = True
except ImportError:
    DOCXCOMPOSE_AVAILABLE = False

logger = logging.getLogger(__name__)


def generer_document_word(data, args, output_path):
    """
    Génère le document Word CISR complet.

    Args:
        data: TranscriptionData avec paragraphes, metadata, etc.
        args: Arguments CLI (section, dossier, iuc, huis_clos, etc.)
        output_path: Chemin de sortie du fichier .docx

    Returns:
        Path du fichier généré
    """
    try:
        logger.info("Génération document Word...")

        doc = Document()

        # Style Normal = Arial 11pt
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Marges CISR section 1 (en-tête)
        margins = _get_margins(args.section)
        section = doc.sections[0]
        section.top_margin = Inches(margins['top'])
        section.bottom_margin = Inches(margins['bottom'])
        section.left_margin = Inches(margins['left'])
        section.right_margin = Inches(margins['right'])

        # En-tête CISR
        _ajouter_entete(doc, args)

        # Titre + sous-titre
        _ajouter_titre(doc, args)

        # Saut de section pour marges corps texte
        _ajouter_saut_section_corps(doc)

        # Corps transcription
        _ajouter_corps(doc, data)

        # Tableaux métadonnées (si metadata JSON disponible)
        if hasattr(args, 'metadata_json') and args.metadata_json and os.path.exists(args.metadata_json):
            _ajouter_tableaux_metadata(doc, args.metadata_json)

        # Marqueur FIN + Certification
        _ajouter_fin_et_certification(doc, args)

        # Fusion avec page couverture si fournie
        output_path = _gerer_page_couverture(doc, args, output_path)

        return output_path

    except Exception as e:
        raise WorkflowError(f"Échec génération Word : {e}") from e


def _get_margins(section_type: str) -> dict:
    """Retourne les marges selon le type de document."""
    margins_map = {
        'SPR': MARGINS_SPR,
        'SAR': MARGINS_SAR,
        'SI': MARGINS_SI,
        'SAI': MARGINS_SAI,
    }
    return margins_map.get(section_type, MARGINS_SPR)


def _ajouter_entete(doc, args):
    """Ajoute l'en-tête CISR bilingue (dossier, IUC, huis clos)."""
    # Dossier / File
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Dossier de la {args.section}")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run = p.add_run(f" / RPD File: {args.dossier}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    # IUC
    if args.iuc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("IUC / UCI:")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        iuc_valeurs = args.iuc.replace(',', '\n').strip()
        run = p.add_run(f" {iuc_valeurs}")
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    doc.add_paragraph()  # Ligne vierge

    # Huis clos
    if args.huis_clos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Huis clos / Private Proceeding")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()


def _ajouter_titre(doc, args):
    """Ajoute le titre et sous-titre CISR."""
    # Titre
    p_titre = doc.add_paragraph()
    run_titre = p_titre.add_run("TRANSCRIPTION DES Motifs et de la décision")
    run_titre.font.name = 'Arial'
    run_titre.font.size = Pt(11)
    p_titre.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # Sous-titre centré, gras, souligné
    p = doc.add_paragraph()
    run = p.add_run("décision")
    run.bold = True
    run.underline = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def _ajouter_saut_section_corps(doc):
    """Ajoute un saut de section continu avec marges standard pour le corps."""
    from docx.enum.section import WD_SECTION

    doc.add_paragraph()
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    new_section.top_margin = Inches(0.89)
    new_section.bottom_margin = Inches(1.00)
    new_section.left_margin = Inches(1.00)
    new_section.right_margin = Inches(1.00)


def _ajouter_corps(doc, data):
    """Ajoute le corps de la transcription avec formatage locuteurs."""
    for i, paragraphe in enumerate(data.paragraphes):
        est_titre, nom_titre = detecter_titre_section(paragraphe)

        p = doc.add_paragraph(paragraphe)
        p.style = 'Normal'
        p.paragraph_format.line_spacing = 1.0

        # Indentation première ligne pour locuteurs
        if paragraphe.strip().startswith(('COMMISSAIRE :', 'CONSEIL :', 'DEMANDEUR :', 'INTERPRÈTE :')):
            p.paragraph_format.first_line_indent = Inches(0.15)

        # Formatage runs
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            # Premier paragraphe COMMISSAIRE en BOLD
            if i == 0 and paragraphe.startswith("COMMISSAIRE :"):
                run.bold = True

            # Titres de sections en GRAS
            if est_titre:
                run.bold = True

        # Ligne vierge après (sauf dernier)
        if i < len(data.paragraphes) - 1:
            doc.add_paragraph()


def _ajouter_tableaux_metadata(doc, metadata_json_path):
    """Ajoute les tableaux de métadonnées bilingues CISR."""
    logger.info("Ajout des tableaux métadonnées CISR...")

    with open(metadata_json_path, 'r', encoding='utf-8') as f:
        metadata = json.load(f)

    doc.add_paragraph()

    # TABLEAU 1: 17L x 3C (9 données + 8 lignes vides intercalaires)
    table = doc.add_table(rows=17, cols=3)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(2.31)
    table.columns[1].width = Inches(3.00)
    table.columns[2].width = Inches(2.44)

    rows_data = [
        ("Demandeur(e)(s) d'asile", metadata.get('participants', {}).get('demandeur', ''), "Claimant(s)"),
        ("Date de l'audience", metadata.get('audience', {}).get('date', ''), "Date of hearing"),
        ("Lieu de l'audience", metadata.get('audience', {}).get('lieu', ''), "Place of hearing"),
        ("Date de la décision / et des motifs", metadata.get('audience', {}).get('date_decision', ''), "Date of decision / and reasons"),
        ("Tribunal", metadata.get('participants', {}).get('commissaire', ''), "Panel"),
        ("Conseil(s) du (de la/des) / demandeur(e)(s) d'asile", metadata.get('participants', {}).get('conseil_demandeur', ''), "Counsel(s) for the claimant(s)"),
        ("Représentant(e) désigné(e)", metadata.get('participants', {}).get('representant_designe', ''), "Designated representative"),
        ("Conseil du (de la) ministre", metadata.get('participants', {}).get('conseil_ministre', ''), "Counsel for the Minister"),
        ("Interprète", metadata.get('participants', {}).get('interprete', ''), "Interpreter"),
    ]

    row_idx = 0
    for i, (fr, value, en) in enumerate(rows_data):
        table.rows[row_idx].cells[0].text = fr
        table.rows[row_idx].cells[1].text = str(value)
        table.rows[row_idx].cells[2].text = en

        # Formatage : Col 1 LEFT+GRAS, Col 2 CENTER, Col 3 RIGHT+GRAS
        for j, cell in enumerate(table.rows[row_idx].cells):
            for paragraph in cell.paragraphs:
                if j == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif j == 1:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j == 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    if j in [0, 2]:
                        run.bold = True

        row_idx += 1
        if i < len(rows_data) - 1:
            row_idx += 1  # Ligne vide intercalaire

    # TABLEAU 2: Blanc (séparateur)
    doc.add_paragraph()
    table2 = doc.add_table(rows=1, cols=1)
    table2.style = 'Table Grid'


def _ajouter_fin_et_certification(doc, args):
    """Ajoute le marqueur FIN DES MOTIFS et la certification."""
    doc.add_paragraph()

    # Marqueur FIN
    p = doc.add_paragraph()
    p.add_run("----------FIN DES MOTIFS ----------")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Certification
    p = doc.add_paragraph()
    p.add_run(f"Je, {args.transcripteur}, déclare que cette transcription est exacte.")
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    date_cert = datetime.now().strftime("%d %B %Y")
    p.add_run(date_cert)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run(f"Agence de transcription: {args.agence}")
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)


def _gerer_page_couverture(doc, args, output_path):
    """Fusionne la page couverture client avec le contenu si fournie."""
    page_couverture_path = getattr(args, 'page_couverture', None)

    if not page_couverture_path or not os.path.exists(page_couverture_path):
        doc.save(output_path)
        logger.info(f"Document Word généré (sans page couverture): {output_path}")
        return output_path

    # Sauvegarder contenu dans fichier temporaire
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        contenu_temp_path = tmp.name
    doc.save(contenu_temp_path)

    # SI: Remplir métadonnées page couverture vierge
    page_couverture_remplie = None
    if args.section == 'SI':
        logger.info("Type SI détecté: remplissage métadonnées page couverture...")
        metadata_for_si = {}
        if hasattr(args, 'metadata_json') and args.metadata_json and os.path.exists(args.metadata_json):
            with open(args.metadata_json, 'r', encoding='utf-8') as f:
                metadata_for_si = json.load(f)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_couv:
            page_couverture_remplie = tmp_couv.name
        remplir_metadonnees_si(page_couverture_path, metadata_for_si, page_couverture_remplie)
        page_couverture_path = page_couverture_remplie

    # Fusionner
    logger.info("Fusion page couverture client avec contenu transcription...")
    fusionner_page_couverture_et_contenu(page_couverture_path, contenu_temp_path, output_path)

    # Nettoyer
    try:
        os.unlink(contenu_temp_path)
        if page_couverture_remplie:
            os.unlink(page_couverture_remplie)
    except Exception:
        pass

    logger.info(f"Document Word généré avec page couverture: {output_path}")
    return output_path


def remplir_metadonnees_si(page_couverture_path, metadata, output_path):
    """
    Remplit les métadonnées dans la page couverture SI (template vierge).

    Args:
        page_couverture_path: Chemin page couverture template
        metadata: Dict métadonnées work order
        output_path: Chemin sortie page couverture remplie
    """
    try:
        doc = Document(page_couverture_path)

        # Parcourir les tableaux pour remplir les champs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texte_cell = cell.text.strip().upper()

                    # Matching des champs par label
                    if 'DOSSIER' in texte_cell or 'FILE NO' in texte_cell:
                        _remplir_cellule_adjacente(row, cell, metadata.get('numero_dossier', ''))
                    elif 'DATE' in texte_cell and 'AUDIENCE' in texte_cell:
                        _remplir_cellule_adjacente(row, cell, metadata.get('date_audience', ''))
                    elif 'PERSONNE CONCERNÉE' in texte_cell or 'CONCERNED PERSON' in texte_cell:
                        _remplir_cellule_adjacente(row, cell, metadata.get('demandeur', ''))
                    elif 'COMMISSAIRE' in texte_cell or 'MEMBER' in texte_cell:
                        _remplir_cellule_adjacente(row, cell, metadata.get('commissaire', ''))

        doc.save(output_path)
        logger.info(f"Page couverture SI remplie: {output_path}")

    except Exception as e:
        raise WorkflowError(f"Échec remplissage page couverture SI: {e}") from e


def _remplir_cellule_adjacente(row, cell_label, valeur):
    """Remplit la cellule adjacente à un label dans un tableau Word."""
    cells = list(row.cells)
    for i, c in enumerate(cells):
        if c == cell_label and i + 1 < len(cells):
            cells[i + 1].text = str(valeur)
            break


def fusionner_page_couverture_et_contenu(page_couverture_path, contenu_path, output_path):
    """
    Fusionne page couverture + contenu transcription.

    Utilise docxcompose si disponible, sinon fusion manuelle python-docx.

    Args:
        page_couverture_path: Chemin page couverture .docx
        contenu_path: Chemin contenu transcription .docx
        output_path: Chemin document fusionné .docx
    """
    try:
        if DOCXCOMPOSE_AVAILABLE:
            logger.info("Fusion via docxcompose...")
            doc_couverture = Document(page_couverture_path)
            doc_contenu = Document(contenu_path)
            composer = Composer(doc_couverture)
            composer.append(doc_contenu)
            composer.save(output_path)
        else:
            logger.info("docxcompose non disponible, fusion manuelle python-docx...")
            doc_couverture = Document(page_couverture_path)
            doc_contenu = Document(contenu_path)

            doc_couverture.add_page_break()
            for element in doc_contenu.element.body:
                new_element = copy.deepcopy(element)
                doc_couverture.element.body.append(new_element)

            doc_couverture.save(output_path)

        logger.info(f"Document fusionné créé: {output_path}")
        return output_path

    except Exception as e:
        raise WorkflowError(f"Échec fusion documents: {e}") from e
