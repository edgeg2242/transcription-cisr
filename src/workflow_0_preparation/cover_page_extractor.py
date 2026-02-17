"""
Extraction des metadonnees depuis les pages couvertures CISR.

Supporte les types SPR et SAR avec leurs structures distinctes:
- SPR: Tableau 9 lignes x 3 colonnes (bilingue FR/EN) avec lignes vides intercalaires
- SAR: Tableau simplifie 1x2 (numeros dossiers seulement), multi-cas
"""
import os
import re
import logging
from docx import Document
from src.common.exceptions import WorkOrderError

logger = logging.getLogger(__name__)


def extraire_metadonnees_spr(page_couverture_path):
    """
    Extraire metadonnees depuis page couverture SPR.

    Structure attendue:
    - En-tete: Numero dossier, IUC, Huis clos
    - Tableau 1 (9 lignes x 3 colonnes): Metadonnees bilingues FR/EN
      avec lignes vides intercalaires (17 lignes physiques)

    Args:
        page_couverture_path: Chemin vers fichier Word page couverture

    Returns:
        dict: Metadonnees extraites avec cles:
            numero_dossier, iuc, huis_clos, section,
            demandeur, date_audience, lieu_audience, date_decision,
            commissaire, conseil_demandeur, representant_designe,
            conseil_ministre, interprete

    Raises:
        WorkOrderError: Si fichier invalide ou structure incorrecte
    """
    try:
        doc = Document(page_couverture_path)

        # --- EXTRACTION EN-TETE ---

        if len(doc.paragraphs) < 2:
            raise WorkOrderError("Page couverture: pas assez de paragraphes")

        # Paragraphe 1: "Dossier de la SPR / RPD File: TC5-07390"
        para_dossier = doc.paragraphs[0].text
        match_dossier = re.search(r'[A-Z]{2}\d-\d+', para_dossier)
        if not match_dossier:
            raise WorkOrderError(f"Numero dossier non trouve dans: {para_dossier}")
        numero_dossier = match_dossier.group()

        # Paragraphe 2: "IUC / UCI: 1118522122"
        para_iuc = doc.paragraphs[1].text
        match_iuc = re.search(r'\d{10}', para_iuc)
        iuc = match_iuc.group() if match_iuc else None

        # Chercher "Huis clos" dans les 5 premiers paragraphes
        huis_clos = False
        for i in range(min(5, len(doc.paragraphs))):
            if "Huis clos" in doc.paragraphs[i].text or "Private Proceeding" in doc.paragraphs[i].text:
                huis_clos = True
                break

        # --- EXTRACTION TABLEAU METADONNEES ---

        if len(doc.tables) == 0:
            raise WorkOrderError("Page couverture: aucun tableau trouve")

        table = doc.tables[0]  # Premier tableau

        # Extraire lignes non-vides du tableau (ignorer lignes vides intercalaires)
        # Les tableaux CISR contiennent des lignes vides entre chaque ligne de donnees
        lignes_donnees = []
        for row in table.rows:
            # Verifier si la ligne contient des donnees (cellule du milieu non vide)
            texte_milieu = row.cells[1].text.strip()
            if texte_milieu:
                lignes_donnees.append(row)

        if len(lignes_donnees) < 9:
            raise WorkOrderError(
                f"Page couverture: pas assez de lignes de donnees "
                f"({len(lignes_donnees)}, attendu 9)"
            )

        # Mapping lignes de donnees (index 0-8)
        metadonnees = {
            'numero_dossier': numero_dossier,
            'iuc': iuc,
            'huis_clos': huis_clos,
            'section': 'SPR',
            'demandeur': lignes_donnees[0].cells[1].text.strip(),
            'date_audience': lignes_donnees[1].cells[1].text.strip(),
            'lieu_audience': lignes_donnees[2].cells[1].text.strip(),
            'date_decision': lignes_donnees[3].cells[1].text.strip(),
            'commissaire': lignes_donnees[4].cells[1].text.strip(),
            'conseil_demandeur': lignes_donnees[5].cells[1].text.strip(),
            'representant_designe': lignes_donnees[6].cells[1].text.strip(),
            'conseil_ministre': lignes_donnees[7].cells[1].text.strip(),
            'interprete': lignes_donnees[8].cells[1].text.strip()
        }

        logger.info(f"Metadonnees SPR extraites: {len(metadonnees)} champs")

        return metadonnees

    except WorkOrderError:
        raise
    except Exception as e:
        raise WorkOrderError(f"Erreur extraction metadonnees SPR: {e}")


def extraire_metadonnees_sar(extracted_dir):
    """
    Extraire metadonnees SAR (multi-dossiers).

    Les pages couvertures SAR ont une structure simplifiee:
    - Tableau 1x2 avec numeros de dossiers (MC5-xxxxx SAR, MC2/3-xxxxx SPR)
    - Pas de metadonnees completes (demandeur, commissaire, etc.)

    Etapes:
    1. Parser toutes les pages couvertures .docx
    2. Scanner dossiers audio recursivement
    3. Fusionner donnees couverture + audio
    4. Validation croisee

    Args:
        extracted_dir: Dossier extrait du ZIP

    Returns:
        dict: Metadonnees SAR avec structure:
            {
                'work_order': {},
                'cases': [
                    {
                        'case_id': 'MC5-xxxxx',
                        'spr_file': 'MC3-xxxxx',
                        'uci': '1234567890',
                        'cover_page_file': '...',
                        'audio_folder': 'MC3-xxxxx',
                        'audio_data': {...}
                    },
                    ...
                ]
            }

    Raises:
        WorkOrderError: Si donnees manquantes ou incoherentes
    """
    metadata = {
        'work_order': {},
        'cases': []
    }

    # === 1. PARSER PAGES COUVERTURES ===

    cover_pages = {}
    docx_files = []
    for root, dirs, files in os.walk(extracted_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                docx_files.append(os.path.join(root, file))

    logger.info(f"Pages couvertures trouvees: {len(docx_files)}")

    for docx_file in docx_files:
        if ("Page couverture" in docx_file
                or "Transcript" in docx_file
                or "MC5-" in os.path.basename(docx_file)):
            try:
                doc = Document(docx_file)

                # Extraire numeros SAR et SPR depuis:
                # 1. Nom du fichier (ex: MC5-34977 Irb 101.41 Page couverture...)
                # 2. En-tete du document (paragraphes 0-1)
                filename = os.path.basename(docx_file)
                para_0 = doc.paragraphs[0].text if len(doc.paragraphs) > 0 else ""
                para_1 = doc.paragraphs[1].text if len(doc.paragraphs) > 1 else ""

                # SAR number (MC5-xxxxx) - chercher dans nom fichier ET contenu
                match_sar = re.search(r'MC5-\d+', filename + para_0 + para_1)
                if not match_sar:
                    logger.warning(f"Numero SAR non trouve dans {filename}, ignore")
                    continue
                sar_number = match_sar.group()

                # Extraire metadonnees depuis tableau
                # NOTE: Pages couvertures SAR ont structure simplifiee
                # (1 ligne, infos de base seulement)
                claimant = ""
                hearing_dates = []
                panel = ""
                counsel = ""
                uci = None
                spr_number = None

                if len(doc.tables) > 0:
                    table = doc.tables[0]

                    # Extraire SPR number et IUC depuis tableau
                    for row in table.rows:
                        for cell in row.cells:
                            # SPR number (MC2-xxxxx ou MC3-xxxxx)
                            if not spr_number:
                                match_spr = re.search(r'MC[2-3]-\d+', cell.text)
                                if match_spr:
                                    spr_number = match_spr.group()

                            # IUC
                            if not uci:
                                match_uci = re.search(r'\d{10}', cell.text)
                                if match_uci:
                                    uci = match_uci.group()

                # Fallback: chercher dans paragraphes
                if not spr_number:
                    match_spr = re.search(r'MC[2-3]-\d+', para_0 + para_1)
                    spr_number = match_spr.group() if match_spr else None

                # Extraire IUC depuis paragraphes si non trouve dans tableau
                if not uci:
                    for p in doc.paragraphs[:10]:
                        match_uci = re.search(r'\d{10}', p.text)
                        if match_uci:
                            uci = match_uci.group()
                            break

                cover_pages[sar_number] = {
                    'spr_file': spr_number,
                    'claimant': claimant,
                    'hearing_dates': hearing_dates,
                    'panel': panel,
                    'counsel': counsel,
                    'uci': uci,
                    'cover_page_file': os.path.basename(docx_file)
                }

                logger.info(
                    f"  - {sar_number}: {claimant}, "
                    f"{len(hearing_dates)} audience(s)"
                )

            except Exception as e:
                logger.warning(f"Erreur parsing {os.path.basename(docx_file)}: {e}")
                continue

    if not cover_pages:
        raise WorkOrderError("SAR: Aucune page couverture valide trouvee")

    # === 2. SCANNER DOSSIERS AUDIO ===

    audio_mappings = {}
    for item in os.listdir(extracted_dir):
        item_path = os.path.join(extracted_dir, item)
        if os.path.isdir(item_path) and item.startswith("MC"):
            folder_name = item  # MC2-27593
            daudio_path = os.path.join(item_path, "DAUDIO")

            if os.path.exists(daudio_path):
                dates_times = {}
                for date_dir in os.listdir(daudio_path):
                    date_path = os.path.join(daudio_path, date_dir)
                    if os.path.isdir(date_path):
                        # Chercher fichiers audio dans ce dossier et sous-dossiers
                        audio_files = []
                        for root, dirs, files in os.walk(date_path):
                            for file in files:
                                if file.endswith(('.a00', '.wav', '.mp3')):
                                    audio_files.append(file)
                        if audio_files:
                            dates_times[date_dir] = audio_files

                if dates_times:
                    audio_mappings[folder_name] = dates_times
                    logger.info(
                        f"  Audio: {folder_name} -> {len(dates_times)} dates"
                    )

    # === 3. FUSIONNER DONNEES ===

    for sar_id, cover_data in cover_pages.items():
        spr_id = cover_data['spr_file']

        case = {
            'case_id': sar_id,
            'spr_file': spr_id,
            'claimant': cover_data['claimant'],
            'panel': cover_data['panel'],
            'counsel': cover_data['counsel'],
            'uci': cover_data['uci'],
            'hearing_dates': cover_data['hearing_dates'],
            'cover_page_file': cover_data['cover_page_file'],
            'audio_folder': spr_id,
            'audio_data': audio_mappings.get(spr_id, {})
        }

        metadata['cases'].append(case)

    # === 4. VALIDATION ===

    for case in metadata['cases']:
        if not case.get('claimant'):
            logger.warning(f"SAR {case['case_id']}: Demandeur manquant")

        if not case['audio_data']:
            logger.warning(
                f"SAR {case['case_id']}: Aucun dossier audio trouve "
                f"pour SPR {case['spr_file']}"
            )

    logger.info(f"{len(metadata['cases'])} cas SAR identifies")

    return metadata


def parser_dates_multiples(texte_dates):
    """
    Parse dates multiples depuis texte francais.

    Exemples:
    - "4 juin 2025" -> ['2025-06-04']
    - "4 juin 2025 et le 19 septembre 2025" -> ['2025-06-04', '2025-09-19']

    Args:
        texte_dates: Texte contenant dates francaises

    Returns:
        list: Liste de dates au format YYYY-MM-DD
    """
    # Dictionnaire mois francais -> numero
    mois_fr = {
        'janvier': 1, 'fevrier': 2, 'mars': 3, 'avril': 4,
        'mai': 5, 'juin': 6, 'juillet': 7, 'aout': 8,
        'septembre': 9, 'octobre': 10, 'novembre': 11, 'decembre': 12,
        # Avec accents
        'f\u00e9vrier': 2, 'ao\u00fbt': 8, 'd\u00e9cembre': 12,
        # Abbreviations
        'janv': 1, 'f\u00e9vr': 2, 'avr': 4,
        'juil': 7, 'sept': 9, 'oct': 10, 'nov': 11, 'd\u00e9c': 12
    }

    # Regex pour dates francaises
    pattern = r'(\d{1,2})\s+([a-z\u00fb\u00e9\u00e8]+)\s+(\d{4})'
    matches = re.findall(pattern, texte_dates, re.IGNORECASE)

    dates = []
    for jour, mois_str, annee in matches:
        mois_str = mois_str.lower()
        if mois_str in mois_fr:
            mois = mois_fr[mois_str]
            date_str = f"{annee}-{mois:02d}-{int(jour):02d}"
            dates.append(date_str)
        else:
            logger.warning(f"Mois non reconnu: {mois_str}")

    return dates
